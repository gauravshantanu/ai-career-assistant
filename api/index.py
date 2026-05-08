"""
AI Career Assistant — Flask version
Run: python3 flask_app.py
Open: http://127.0.0.1:5000
"""

from flask import Flask, render_template_string, request, session, jsonify, send_file
import os, io, json, urllib.parse
from datetime import datetime, timedelta
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor

# ── Load environment variables from .env file ─────────────────────────────────
def load_env():
    # Load .env for local dev; on Vercel env vars are set in dashboard
    try:
        env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
        if os.path.exists(env_path):
            for line in open(env_path):
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k.strip(), v.strip())
    except: pass

load_env()

app = Flask(__name__)
app.secret_key = "career_ai_secret_key_2024_xyz"
app.config["SESSION_PERMANENT"] = True
from datetime import timedelta
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=30)

@app.route("/favicon.ico")
def favicon():
    return "", 204

KEY_FILE       = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".groq_key")  # local only
ADMIN_PASSWORD = "admin@career2024"

LANGUAGES = {
    "English":"English","हिन्दी (Hindi)":"Hindi","தமிழ் (Tamil)":"Tamil",
    "తెలుగు (Telugu)":"Telugu","বাংলা (Bengali)":"Bengali","मराठी (Marathi)":"Marathi",
    "ಕನ್ನಡ (Kannada)":"Kannada","മലയാളം (Malayalam)":"Malayalam",
    "ગુજરાતી (Gujarati)":"Gujarati","اردو (Urdu)":"Urdu","Español":"Spanish","Français":"French",
}

TOOLS = [
    ("resume",    "description", "Resume Review",      "Upload your resume and get AI-powered feedback, score, and improvements."),
    ("interview", "forum",       "Mock Interview",     "Practice real interview questions with live AI feedback."),
    ("cover",     "mail",        "Cover Letter",       "Generate a polished, human-sounding cover letter instantly."),
    ("linkedin",  "share",       "LinkedIn Post",      "Turn your achievements into viral LinkedIn posts."),
    ("decoder",   "analytics",   "Job Decoder",        "Understand what any job posting really wants."),
    ("apply",     "touch_app",   "Apply One-Click",    "Get a full application package — cover letter, email, talking points."),
    ("jobs",      "search",      "Job Finder",         "AI matches your profile to jobs on Naukri, LinkedIn and more."),
    ("career",    "map",         "Career Path",        "Get a detailed roadmap from where you are to where you want to be."),
    ("salary",    "payments",    "Salary Coach",       "Know your worth, negotiate confidently, write counter offers."),
    ("history",   "history",     "My History",         "All your generated content saved in one place."),
]

def load_key():
    # Priority: .env file → environment variable → .groq_key file
    return (
        os.environ.get("GROQ_API_KEY", "") or
        (open(KEY_FILE).read().strip() if os.path.exists(KEY_FILE) else "")
    )

def save_key(k):
    # On Vercel filesystem is read-only — just update env var in memory
    os.environ["GROQ_API_KEY"] = k
    # Also try writing to .env for local dev
    try:
        env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
        lines = []
        found = False
        if os.path.exists(env_path):
            for line in open(env_path):
                if line.startswith("GROQ_API_KEY="):
                    lines.append(f"GROQ_API_KEY={k}\n")
                    found = True
                else:
                    lines.append(line)
        if not found:
            lines.append(f"GROQ_API_KEY={k}\n")
        open(env_path, "w").writelines(lines)
    except: pass

def call_ai(api_key, system, user, history=None):
    msgs = [{"role":"system","content":system}] + (history or []) + [{"role":"user","content":user}]
    client = Groq(api_key=api_key)
    r = client.chat.completions.create(model="llama-3.3-70b-versatile", max_tokens=1500, messages=msgs)
    return r.choices[0].message.content

def ln(lang):
    return "" if lang == "English" else f"\n\nIMPORTANT: Write ENTIRE response in {lang}."

def mk_docx(title, text):
    doc = Document()
    h = doc.add_heading(title, 0); h.runs[0].font.color.rgb = RGBColor(0x00,0x59,0xb5)
    p = doc.add_paragraph(f"Generated {datetime.now().strftime('%d %B %Y')}"); p.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    for line in text.split("\n"):
        if line.strip(): doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ── Base HTML template ────────────────────────────────────────────────────────
BASE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>CareerAI — {{ tool_name }}</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet"/>
<link href="https://fonts.googleapis.com/css2?family=Material+Symbols+Outlined:opsz,wght,FILL,GRAD@24,400,0,0&display=swap" rel="stylesheet"/>
<style>
:root{
  --bg:#f5f5f7;--white:#ffffff;
  --surf-low:#f2f3fd;--surf:#ecedf7;--surf-high:#e6e8f2;--surf-highest:#e0e2ec;
  --on-surf:#181c23;--on-surf-var:#414753;
  --outline:#717785;--outline-var:#c1c6d6;
  --primary:#0059b5;--primary-c:#0071e3;--on-primary-c:#ffffff;
  --secondary:#5e5e63;--sec-c:#e0dfe4;--on-sec-c:#626267;
  --tertiary:#9b3f00;--tert-c:#c25100;--on-tert-c:#ffffff;
  --error:#ba1a1a;
  --font:Inter,-apple-system,BlinkMacSystemFont,sans-serif;
  --nav-h:52px;--side-w:256px;
  --shadow:0 4px 24px rgba(0,0,0,0.04);
  --shadow-md:0 8px 32px rgba(0,0,0,0.08);
}
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:var(--font);background:var(--bg);color:var(--on-surf);-webkit-font-smoothing:antialiased;}

/* NAV */
.nav{position:fixed;top:0;left:0;right:0;height:var(--nav-h);z-index:999;
  background:rgba(249,249,255,0.85);backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);
  border-bottom:0.5px solid var(--outline-var);
  display:flex;align-items:center;justify-content:space-between;padding:0 24px;}
.nav-brand{font-size:22px;font-weight:700;color:var(--on-surf);letter-spacing:-0.02em;text-decoration:none;}
.nav-links{display:flex;gap:24px;margin-left:32px;}
.nav-link{color:var(--secondary);font-size:15px;font-weight:500;text-decoration:none;transition:color .15s;}
.nav-link:hover{color:var(--primary);}
.nav-right{display:flex;align-items:center;gap:10px;margin-left:auto;}
.nav-icon-btn{background:none;border:none;cursor:pointer;color:var(--on-surf-var);
  font-size:22px;display:flex;align-items:center;transition:color .15s;padding:4px;}
.nav-icon-btn:hover{color:var(--primary);}
.upgrade-btn{background:var(--primary-c);color:var(--on-primary-c);
  padding:5px 16px;border-radius:8px;font-size:13px;font-weight:600;
  border:none;cursor:pointer;text-decoration:none;transition:opacity .15s;}
.upgrade-btn:hover{opacity:.88;}
.avatar{width:32px;height:32px;border-radius:50%;background:var(--surf);
  border:1px solid var(--outline-var);display:flex;align-items:center;
  justify-content:center;font-size:12px;font-weight:600;color:var(--primary);}

/* SIDEBAR */
.sidebar{position:fixed;left:0;top:var(--nav-h);bottom:0;width:var(--side-w);
  background:var(--surf-low);border-right:0.5px solid var(--outline-var);
  overflow-y:auto;padding:16px;z-index:100;display:flex;flex-direction:column;gap:2px;}
.sidebar-hd{padding:8px 4px 16px;}
.sidebar-title{font-size:20px;font-weight:700;color:var(--primary);letter-spacing:-0.01em;}
.sidebar-sub{font-size:12px;font-weight:500;color:var(--secondary);margin-top:2px;letter-spacing:0.01em;}
.nav-item{display:flex;align-items:center;gap:10px;padding:9px 12px;border-radius:8px;
  text-decoration:none;color:var(--on-surf-var);font-size:13px;font-weight:500;
  letter-spacing:0.01em;transition:background .15s,color .15s;}
.nav-item:hover{background:var(--surf-highest);color:var(--on-surf);}
.nav-item.active{background:var(--sec-c);color:var(--on-sec-c);font-weight:700;}
.nav-item .ms{font-size:20px;flex-shrink:0;}

/* MAIN */
.main{margin-left:var(--side-w);padding-top:var(--nav-h);}
.content{max-width:1100px;padding:40px 48px 80px;}

/* PAGE HEADER */
.page-title{font-size:44px;font-weight:700;color:var(--on-surf);
  letter-spacing:-0.022em;line-height:1.1;margin-bottom:8px;}
.page-sub{font-size:17px;color:var(--secondary);letter-spacing:-0.015em;
  line-height:1.5;margin-bottom:32px;}

/* CARDS */
.card{background:var(--white);border-radius:16px;border:1px solid rgba(193,198,214,0.25);
  box-shadow:var(--shadow);padding:24px;margin-bottom:16px;
  transition:box-shadow .2s;}
.card-title{font-size:20px;font-weight:600;color:var(--on-surf);
  letter-spacing:-0.007em;margin-bottom:16px;}

/* GRID */
.grid2{display:grid;grid-template-columns:1fr 1fr;gap:20px;}
.grid3{display:grid;grid-template-columns:1fr 1fr 1fr;gap:16px;}
@media(max-width:768px){
  .grid2,.grid3{grid-template-columns:1fr;}
  .sidebar{display:none;}
  .main{margin-left:0;}
  .content{padding:24px 20px 60px;}
  .nav-links{display:none;}
  .page-title{font-size:28px;}
}
@media(max-width:1024px){
  .sidebar{width:200px;}
  .main{margin-left:200px;}
  .block-container{padding-left:calc(200px + 1rem);}
}

/* FORM ELEMENTS */
.form-group{margin-bottom:16px;}
label{display:block;font-size:13px;font-weight:500;color:var(--secondary);
  letter-spacing:0.01em;margin-bottom:6px;}
input[type=text],input[type=email],input[type=password],textarea,select{
  width:100%;background:var(--white);border:1px solid var(--outline-var);
  border-radius:8px;padding:10px 12px;font-family:var(--font);font-size:15px;
  color:var(--on-surf);letter-spacing:-0.015em;outline:none;
  transition:border .15s,box-shadow .15s;}
select{appearance:none;background-image:url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='16' height='16' viewBox='0 0 24 24' fill='none' stroke='%23717785' stroke-width='2'%3E%3Cpath d='M6 9l6 6 6-6'/%3E%3C/svg%3E");
  background-repeat:no-repeat;background-position:right 12px center;padding-right:36px;}
input:focus,textarea:focus,select:focus{
  border-color:var(--primary-c);box-shadow:0 0 0 3px rgba(0,113,227,0.12);outline:none;}
input::placeholder,textarea::placeholder{color:#a8aab4;font-size:14px;}
textarea{resize:vertical;line-height:1.5;}

/* BUTTONS */
.btn{display:inline-flex;align-items:center;gap:8px;padding:10px 20px;
  border-radius:8px;font-family:var(--font);font-size:15px;font-weight:600;
  cursor:pointer;border:none;text-decoration:none;transition:opacity .15s,transform .1s;
  letter-spacing:-0.01em;}
.btn:active{transform:scale(0.98);}
.btn-primary{background:var(--primary-c);color:var(--on-primary-c);width:100%;
  justify-content:center;font-size:15px;padding:12px;}
.btn-primary:hover{opacity:.9;}
.btn-secondary{background:var(--sec-c);color:var(--on-sec-c);}
.btn-secondary:hover{background:var(--surf-highest);}
.btn-ghost{background:transparent;color:var(--primary);border:1px solid var(--primary-c);}
.btn-danger{background:#ffdad6;color:#ba1a1a;}
.btn-sm{padding:6px 14px;font-size:13px;}
.btn-full{width:100%;justify-content:center;}

/* TAGS / CHIPS */
.chip{display:inline-flex;align-items:center;background:var(--sec-c);
  color:var(--on-sec-c);padding:4px 12px;border-radius:9999px;
  font-size:12px;font-weight:500;margin:3px;letter-spacing:0.01em;}
.chip-active{background:var(--primary-c);color:var(--on-primary-c);}

/* OUTPUT */
.output{background:var(--white);border-radius:16px;border:1px solid rgba(193,198,214,0.25);
  box-shadow:var(--shadow);padding:24px;margin-top:16px;}
.output-text{font-size:15px;line-height:1.7;letter-spacing:-0.01em;
  white-space:pre-wrap;color:var(--on-surf);}
.output-label{font-size:11px;font-weight:600;text-transform:uppercase;
  letter-spacing:0.08em;color:var(--secondary);margin-bottom:12px;}
.action-bar{display:flex;gap:8px;flex-wrap:wrap;margin-top:16px;padding-top:16px;
  border-top:1px solid var(--outline-var);}
.action-bar .btn{font-size:12px;padding:6px 12px;border-radius:6px;}

/* CHAT */
.chat-wrap{display:flex;flex-direction:column;gap:8px;margin:16px 0;
  max-height:420px;overflow-y:auto;padding-right:4px;
  scrollbar-width:thin;scrollbar-color:var(--outline-var) transparent;}
.bubble-ai{background:var(--white);border:0.5px solid var(--outline-var);
  border-radius:12px 12px 12px 4px;box-shadow:var(--shadow);
  padding:14px 18px;max-width:84%;font-size:15px;line-height:1.6;}
.bubble-user{background:var(--primary-c);color:var(--on-primary-c);
  border-radius:12px 12px 4px 12px;
  padding:14px 18px;max-width:84%;margin-left:auto;font-size:15px;line-height:1.6;font-weight:500;}

/* TABS */
.tabs{display:flex;background:var(--surf);border-radius:8px;padding:4px;gap:2px;margin-bottom:20px;}
.tab-btn{flex:1;padding:7px 0;border:none;border-radius:6px;background:transparent;
  color:var(--secondary);font-family:var(--font);font-size:13px;font-weight:500;
  cursor:pointer;letter-spacing:0.01em;transition:all .15s;}
.tab-btn.active{background:var(--primary-c);color:var(--on-primary-c);font-weight:600;
  box-shadow:0 2px 8px rgba(0,0,0,0.08);}
.tab-content{display:none;}
.tab-content.active{display:block;}

/* APPLY LINKS */
.apply-box{background:var(--white);border-radius:16px;border:1px solid rgba(193,198,214,0.25);
  box-shadow:var(--shadow);padding:24px;margin-top:16px;}
.apply-link{display:inline-flex;align-items:center;gap:6px;background:var(--sec-c);
  color:var(--on-sec-c);padding:8px 16px;border-radius:8px;text-decoration:none;
  font-size:13px;font-weight:500;margin:4px 3px;transition:all .15s;}
.apply-link:hover{background:var(--surf-highest);color:var(--on-surf);}

/* UPLOAD ZONE */
.upload-zone{border:2px dashed var(--outline-var);border-radius:12px;padding:32px;
  text-align:center;cursor:pointer;background:rgba(242,243,253,0.5);
  transition:border-color .15s;}
.upload-zone:hover{border-color:var(--primary-c);}
.upload-icon{font-size:40px;color:var(--primary-c);margin-bottom:12px;}

/* ALERTS */
.alert{padding:12px 16px;border-radius:8px;font-size:14px;margin:8px 0;}
.alert-error{background:rgba(186,26,26,0.08);color:#ba1a1a;}
.alert-success{background:rgba(0,89,181,0.07);color:var(--primary);}
.alert-warn{background:rgba(155,63,0,0.07);color:var(--tertiary);}

/* HISTORY */
.history-item{background:var(--white);border-radius:12px;border:1px solid rgba(193,198,214,0.25);
  box-shadow:var(--shadow);margin-bottom:12px;overflow:hidden;}
.history-head{padding:14px 20px;cursor:pointer;display:flex;align-items:center;
  justify-content:space-between;font-size:14px;font-weight:500;
  border-bottom:1px solid var(--outline-var);}
.history-head:hover{background:var(--surf-low);}
.history-body{padding:20px;display:none;}
.history-body.open{display:block;}

/* EMPTY STATE */
.empty{text-align:center;padding:64px 32px;color:var(--secondary);}
.empty-icon{font-size:48px;margin-bottom:16px;}
.empty-title{font-size:18px;font-weight:600;color:var(--on-surf);margin-bottom:8px;}

/* LOADING SPINNER */
.spinner-overlay{display:none;position:fixed;inset:0;background:rgba(255,255,255,0.7);
  z-index:9999;align-items:center;justify-content:center;flex-direction:column;gap:12px;}
.spinner-overlay.show{display:flex;}
.spinner{width:40px;height:40px;border:3px solid var(--outline-var);
  border-top-color:var(--primary-c);border-radius:50%;animation:spin 0.7s linear infinite;}
@keyframes spin{to{transform:rotate(360deg)}}
.spinner-text{font-size:15px;color:var(--secondary);font-weight:500;}

/* SETTINGS PANEL */
.settings-panel{background:var(--white);border-radius:12px;
  border:0.5px solid var(--outline-var);padding:20px;margin-bottom:20px;display:none;}
.settings-panel.open{display:block;}

/* SCORE CIRCLE */
.score-wrap{display:flex;align-items:center;gap:20px;margin-bottom:20px;}
.score-circle{position:relative;width:72px;height:72px;flex-shrink:0;}
.score-circle svg{width:72px;height:72px;transform:rotate(-90deg);}
.score-num{position:absolute;inset:0;display:flex;align-items:center;justify-content:center;
  font-size:20px;font-weight:700;color:var(--primary-c);}

/* FOOTER */
.footer{border-top:0.5px solid var(--outline-var);padding:24px 48px;
  text-align:center;color:var(--secondary);font-size:12px;}
.footer a{color:var(--secondary);text-decoration:underline;margin:0 8px;}

/* ANIMATIONS */
@keyframes fadeUp{from{opacity:0;transform:translateY(10px);}to{opacity:1;transform:translateY(0);}}
.fade-up{animation:fadeUp .35s ease both;}
@keyframes slideIn{from{opacity:0;transform:translateX(20px);}to{opacity:1;transform:translateX(0);}}
@keyframes slideOut{from{opacity:1;transform:translateX(0);}to{opacity:0;transform:translateX(20px);}}
.nav-item{position:relative;}
.nav-item.active::before{content:'';position:absolute;left:0;top:50%;transform:translateY(-50%);
  width:3px;height:60%;background:var(--primary-c);border-radius:0 2px 2px 0;}

/* Settings toggle */
.settings-toggle{display:flex;align-items:center;gap:6px;font-size:13px;
  color:var(--secondary);cursor:pointer;padding:6px 12px;border-radius:8px;
  border:0.5px solid var(--outline-var);background:var(--white);margin-bottom:8px;
  transition:background .15s;}
.settings-toggle:hover{background:var(--surf-low);}
</style>
</head>
<body>

<!-- SPINNER -->
<div class="spinner-overlay" id="spinner">
  <div class="spinner"></div>
  <div class="spinner-text">✨ Generating with AI...</div>
</div>

<!-- NAV -->
<nav class="nav">
  <div style="display:flex;align-items:center;">
    <a class="nav-brand" href="/">CareerAI</a>
    <div class="nav-links">
      <a class="nav-link" href="#">Tools</a>
      <a class="nav-link" href="#">History</a>
      <a class="nav-link" href="#">Community</a>
    </div>
  </div>
  <div class="nav-right">
    <span style="font-size:12px;color:var(--secondary);font-weight:500;display:none;" id="lang-indicator"></span>
    <button class="nav-icon-btn" onclick="toggleSettings()" title="Settings">
      <span class="material-symbols-outlined">settings</span>
    </button>
    <a class="upgrade-btn" href="https://console.groq.com" target="_blank">Get API Key</a>
    <div class="avatar" title="User">U</div>
  </div>
</nav>

<!-- SIDEBAR -->
<aside class="sidebar">
  <div class="sidebar-hd">
    <div class="sidebar-title">Career Suite</div>
    <div class="sidebar-sub">AI-Powered Tools</div>
  </div>
  {% for slug, icon, label, desc in tools %}
  <a class="nav-item {% if active == slug %}active{% endif %}" href="/tool/{{ slug }}">
    <span class="material-symbols-outlined ms">{{ icon }}</span>{{ label }}
  </a>
  {% endfor %}
</aside>

<!-- MAIN -->
<main class="main">
<div class="content fade-up">

<!-- SETTINGS PANEL -->
<div class="settings-panel" id="settings-panel">
  <form method="POST" action="/settings">
    <div class="grid2" style="gap:16px;align-items:end;">
      <div class="form-group" style="margin:0;">
        <label>Language</label>
        <select name="language">
          {% for k,v in languages.items() %}
          <option value="{{ v }}" {% if v == session.get('lang','English') %}selected{% endif %}>{{ k }}</option>
          {% endfor %}
        </select>
      </div>
      <div class="form-group" style="margin:0;">
        <label>Groq API Key</label>
        <div style="display:flex;gap:8px;">
          <input type="password" name="api_key" value="{{ session.get('api_key','') }}" placeholder="gsk_..."/>
          <button type="submit" class="btn btn-secondary btn-sm" style="white-space:nowrap;">Save</button>
        </div>
      </div>
    </div>
  </form>
</div>

{% block content %}{% endblock %}

</div>

<footer class="footer">
  <strong>CareerAI Assistant</strong> &nbsp;·&nbsp; © 2024 Precision Career Tools
  &nbsp;&nbsp;<a href="#">Privacy</a><a href="#">Terms</a><a href="#">Support</a>
</footer>
</main>

<script>
function toggleSettings(){
  var p = document.getElementById('settings-panel');
  p.classList.toggle('open');
}
// Show current language in nav
(function(){
  var sel = document.querySelector('[name="language"]');
  var ind = document.getElementById('lang-indicator');
  if(sel && ind && sel.value && sel.value !== 'English'){
    ind.style.display='inline';
    ind.textContent=sel.value;
  }
  if(sel) sel.addEventListener('change', function(){
    if(ind){ ind.style.display=this.value!=='English'?'inline':'none'; ind.textContent=this.value; }
  });
})();
function showSpinner(btn){
  document.getElementById('spinner').classList.add('show');
  if(btn){ btn.disabled=true; btn._orig=btn.innerHTML; btn.innerHTML='<span style="opacity:.7">Generating...</span>'; }
}
function _showSpinner(){ document.getElementById('spinner').classList.add('show'); }
function hideSpinner(){
  document.getElementById('spinner').classList.remove('show');
  document.querySelectorAll('button[disabled]').forEach(function(b){
    if(b._orig){ b.innerHTML=b._orig; b.disabled=false; }
  });
}

// Show spinner on all form submissions
document.querySelectorAll('form.ai-form').forEach(function(f){
  f.addEventListener('submit', function(){ showSpinner(); });
});

// Init char counters on page load
document.addEventListener('DOMContentLoaded', function(){
  ['res-text','res-jd','cov-jd','cov-about','li-topic','dec-jd',
   'ap-jd','ap-res','jf-res','cp-skills','ns-ach','ce-why'].forEach(function(id){
    addCharCounter(id, id==='dec-jd'||id==='ap-jd'?5000:2000);
  });
});

// Tabs
function switchTab(group, idx){
  document.querySelectorAll('.tab-content[data-group="'+group+'"]').forEach(function(c,i){
    c.classList.toggle('active', i===idx);
  });
  document.querySelectorAll('.tab-btn[data-group="'+group+'"]').forEach(function(b,i){
    b.classList.toggle('active', i===idx);
  });
}

// History accordion
document.querySelectorAll('.history-head').forEach(function(h){
  h.addEventListener('click', function(){
    var body = this.nextElementSibling;
    body.classList.toggle('open');
  });
});
</script>
</body>
</html>"""

# ── Tool page template pieces ─────────────────────────────────────────────────
def render_output(result, title="", download_key="output", gmail_subject="", linkedin=False):
    gdoc = f"https://docs.google.com/document/create?title={urllib.parse.quote(title)}"
    gmail_html = ""
    if gmail_subject:
        gu = f"https://mail.google.com/mail/?view=cm&su={urllib.parse.quote(gmail_subject)}&body={urllib.parse.quote(result[:1500])}"
        gmail_html = f'<a class="btn btn-secondary btn-sm" href="{gu}" target="_blank">✉️ Gmail</a>'
    li_html = ""
    if linkedin:
        lu = f"https://www.linkedin.com/feed/?shareActive=true&text={urllib.parse.quote(result[:700])}"
        li_html = f'<a class="btn btn-secondary btn-sm" href="{lu}" target="_blank">🔗 LinkedIn</a>'
    return f"""
<div class="output">
  <div class="output-label">AI Result</div>
  <div class="output-text">{result.replace(chr(10),'<br>')}</div>
  <div class="action-bar">
    <a class="btn btn-secondary btn-sm" href="/download/{download_key}">⬇ Download .docx</a>
    <a class="btn btn-secondary btn-sm" href="/download-txt/{download_key}">⬇ .txt</a>
    <a class="btn btn-secondary btn-sm" href="{gdoc}" target="_blank">📄 Google Docs</a>
    {gmail_html}{li_html}
  </div>
</div>"""

def page(active_slug, title, subtitle, body_html, session_obj):
    tools_data = [(s,i,l,d) for s,i,l,d in TOOLS]
    # Build full template: inject body into placeholder, inject shared JS before </body>
    full = (BASE
        .replace("{% block content %}{% endblock %}", body_html)
        .replace("</body>", SHARED_JS + "\n</body>"))
    return render_template_string(
        full,
        tools=tools_data, active=active_slug,
        tool_name=title, languages=LANGUAGES, session=session_obj
    )

# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    # Auto-load saved API key into session if not already set
    if "api_key" not in session:
        session["api_key"] = load_key()
    if "lang" not in session:
        session["lang"] = "English"
    return tool_page("resume")

@app.before_request
def load_defaults():
    if "api_key" not in session:
        session["api_key"] = load_key()
    if "lang" not in session:
        session["lang"] = "English"

@app.route("/tool/<slug>")
def tool_page(slug):
    valid = [t[0] for t in TOOLS]
    if slug not in valid: slug = "resume"
    tool_info = next(t for t in TOOLS if t[0] == slug)
    _, _, label, desc = tool_info

    body = f"""
<div class="page-title">{label}</div>
<div class="page-sub">{desc}</div>
{TOOL_BODIES.get(slug, '<p>Coming soon.</p>')}
"""
    return page(slug, label, desc, body, session)

@app.route("/settings", methods=["POST"])
def settings():
    lang = request.form.get("language","English")
    key  = request.form.get("api_key","").strip()
    session.permanent = True
    session["lang"] = lang
    if key:
        session["api_key"] = key
        save_key(key)
        session.modified = True
    from flask import redirect
    ref = request.referrer or "/"
    from urllib.parse import urlparse
    parsed = urlparse(ref)
    safe = parsed.path + ("?" + parsed.query if parsed.query else "")
    return redirect(safe or "/")

@app.route("/api/<tool>", methods=["POST"])
def api_call(tool):
    key = session.get("api_key","").strip() or load_key().strip()
    if not key:
        return jsonify({"error": "⚠️ No API key found. Click ⚙️ Settings (top right), paste your Groq key and click Save."}), 400
    lang = session.get("lang","English")
    data = request.json or {}
    try:
        result = handle_tool(tool, data, key, lang)
        # Save to history
        hist = session.get("history", [])
        hist.insert(0, {"tool": tool, "time": datetime.now().strftime("%d %b %Y, %I:%M %p"),
                        "summary": str(data.get("summary",""))[:100], "output": result})
        session["history"] = hist[:50]
        # Store for download
        session[f"last_{tool}"] = result
        return jsonify({"result": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/download/<key>")
def download_docx(key):
    text = session.get(f"last_{key}", "No content")
    buf = mk_docx(key.replace("_"," ").title(), text)
    return send_file(buf, as_attachment=True, download_name=f"{key}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/download-txt/<key>")
def download_txt(key):
    text = session.get(f"last_{key}", "No content")
    return send_file(io.BytesIO(text.encode()), as_attachment=True,
                     download_name=f"{key}.txt", mimetype="text/plain")

@app.route("/api/parse-file", methods=["POST"])
def parse_file():
    """Parse uploaded PDF or DOCX and return text"""
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400
    fname = f.filename.lower()
    try:
        if fname.endswith(".pdf"):
            import pdfplumber, io as _io
            with pdfplumber.open(_io.BytesIO(f.read())) as pdf:
                text = chr(10).join([p.extract_text() or "" for p in pdf.pages])
        elif fname.endswith(".docx"):
            from docx import Document as D2
            import io as _io
            doc = D2(_io.BytesIO(f.read()))
            text = chr(10).join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            text = f.read().decode("utf-8", errors="ignore")
        return jsonify({"text": text[:8000]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/history/clear", methods=["POST"])
def clear_history():
    session["history"] = []
    from flask import redirect
    return redirect("/tool/history")

def handle_tool(tool, data, key, lang):
    L = ln(lang)
    if tool == "resume":
        return call_ai(key,
            f"Expert resume coach. Review resume vs job description.\n1. MATCH SCORE (/10)\n2. TOP 3 IMPROVED BULLET POINTS\n3. SKILL GAPS (2-3 missing skills)\n4. QUICK WINS (2 actionable tips)\nBe specific and direct.{L}",
            f"RESUME:\n{data.get('resume','')}\n\nJOB DESCRIPTION:\n{data.get('jd','')}")
    elif tool == "interview_answer":
        return call_ai(key,
            f"Senior interviewer for {data.get('level','')} {data.get('role','')} {data.get('type','')} interview. Give 1-2 sentence feedback on the answer, then ask the next question. After 5 questions give overall score out of 10.{L}",
            data.get("answer",""), history=data.get("history",[]))
    elif tool == "cover":
        return call_ai(key,
            f"Write a compelling {data.get('tone','professional')} cover letter. 3-4 paragraphs. Never start with 'I am writing to'. Be human and specific to the job.{L}",
            f"Candidate:{data.get('name','')}\nEmail:{data.get('email','')}\nCompany:{data.get('company','the company')}\nJob:\n{data.get('jd','')}\nAbout me:\n{data.get('about','')}")
    elif tool == "linkedin":
        return call_ai(key,
            f"Write a compelling LinkedIn post. Type: {data.get('ptype','Career milestone')}. Style: {data.get('style','Storytelling')}. Add 3-5 relevant hashtags at the end. Max 300 words.{L}",
            f"Topic: {data.get('topic','')}")
    elif tool == "decoder":
        return call_ai(key,
            f"Decode this job posting:\n1. WHAT THEY REALLY WANT (3 points)\n2. MUST-HAVE vs NICE-TO-HAVE skills\n3. RED FLAGS (max 3)\n4. WHAT TO HIGHLIGHT in your application\n5. SALARY ESTIMATE (INR range)\n6. 3 SMART QUESTIONS to ask in interview{L}",
            f"Job posting:\n{data.get('jd','')}")
    elif tool == "apply":
        jd, resume = data.get("jd",""), data.get("resume","")
        cov = call_ai(key, f"Write a 3-paragraph cover letter. Never start with 'I am writing to'. Be specific to this exact job.{L}",
                      f"Candidate:{data.get('name','')}\nJob:\n{jd}\nResume:\n{resume}")
        em  = call_ai(key, f"Write a short application email body (3-4 sentences) to send with cover letter attached.{L}",
                      f"Candidate:{data.get('name','')}\nJob:\n{jd[:500]}")
        pts = call_ai(key, f"Give exactly 5 bullet points: the strongest things this candidate should emphasise for this job.{L}",
                      f"Job:\n{jd}\nCandidate:\n{resume}")
        return json.dumps({"cover": cov, "email": em, "points": pts})
    elif tool == "jobs":
        return call_ai(key,
            f"Analyze this profile and give: 1) Top 5 skills to highlight, 2) Best 5 job titles to search, 3) Top 10 companies to target. Be specific and actionable.{L}",
            f"Resume/Role: {data.get('resume', data.get('role',''))[:1000]}\nLocation: {data.get('location','India')}\nExperience: {data.get('exp','')}")
    elif tool == "career":
        return call_ai(key,
            f"Senior career counselor. Create a detailed roadmap:\n🎯 WHERE YOU ARE NOW\n🚀 NEXT STEP (6-12 months): title, salary (INR), skills, certs, actions\n📈 MID-TERM (1-3 years): title, salary, skills, certs\n🏆 LONG-TERM (3-5 years): vision and strategy{L}",
            f"Role: {data.get('role','')}\nExperience: {data.get('exp','')}\nSkills: {data.get('skills','')}\nGoal: {data.get('goal','Not specified')}\nIndustry: {data.get('industry','Technology')}\nTimeline: {data.get('timeline','2 years')}")
    elif tool == "salary_worth":
        return call_ai(key,
            f"Salary expert for Indian job market. Give:\n💰 SALARY RANGE (min/mid/max in LPA)\n📊 MARKET BENCHMARK vs peers\n🏆 TOP PAYING COMPANIES for this role\n📈 FACTORS that increase salary\n💡 NEGOTIATION LEVERAGE points\nUse specific INR/LPA numbers.{L}",
            f"Role: {data.get('role','')}\nExperience: {data.get('exp','')}\nCity: {data.get('city','')}\nSkills: {data.get('skills','')}\nCompany: {data.get('company','General market')}")
    elif tool == "salary_script":
        return call_ai(key,
            f"Write a confident salary negotiation script:\n📞 OPENING LINE\n💪 VALUE STATEMENT\n💬 EXACT WORDS TO SAY\n🔄 HANDLING OBJECTIONS (2-3)\n✅ CLOSING\nMake it natural, confident, not aggressive.{L}",
            f"Offer: {data.get('offer','')} LPA\nExpecting: {data.get('expect','')} LPA\nRole: {data.get('role','')}\nExperience: {data.get('exp','')}\nAchievements: {data.get('achievements','')}")
    elif tool == "salary_email":
        return call_ai(key,
            f"Write a professional counter offer email. Polite but firm. Show value. 3-4 paragraphs.{L}",
            f"Name: {data.get('name','')}\nCompany: {data.get('company','')}\nRole: {data.get('role','')}\nOffer: {data.get('offer','')} LPA\nCounter: {data.get('counter','')} LPA\nReason: {data.get('reason','')}")
    return "Tool not found"

# Inject shared JS helpers into every page
SHARED_JS = """
<script>
/* ── Toast notifications ── */
function toast(msg, type){
  var t=document.createElement('div');
  var bg=type==='error'?'#ba1a1a':type==='warn'?'#9b3f00':'#0059b5';
  t.style.cssText='position:fixed;bottom:24px;right:24px;z-index:99999;background:'+bg+';color:#fff;'+
    'padding:12px 20px;border-radius:8px;font-size:14px;font-weight:500;font-family:Inter,sans-serif;'+
    'box-shadow:0 8px 32px rgba(0,0,0,0.18);animation:slideIn .25s ease;max-width:360px;line-height:1.4;';
  t.textContent=msg;
  document.body.appendChild(t);
  setTimeout(function(){ t.style.animation='slideOut .25s ease forwards'; setTimeout(function(){t.remove()},250); },3500);
}

/* ── Char counter ── */
function addCharCounter(textareaId, max){
  var ta=document.getElementById(textareaId); if(!ta) return;
  var c=document.createElement('div');
  c.style.cssText='font-size:11px;color:#717785;text-align:right;margin-top:3px;';
  c.id=textareaId+'-counter';
  ta.parentNode.insertBefore(c, ta.nextSibling);
  function update(){ c.textContent=ta.value.length+(max?'/'+max:'')+' chars'; }
  ta.addEventListener('input', update); update();
}

/* ── Keyboard shortcut: Ctrl+Enter to submit form ── */
document.addEventListener('keydown', function(e){
  if((e.ctrlKey||e.metaKey) && e.key==='Enter'){
    var focused=document.activeElement;
    if(focused && focused.tagName==='TEXTAREA'){
      var card=focused.closest('.card');
      if(card){
        var btn=card.querySelector('.btn-primary');
        if(btn && !btn.disabled){ btn.click(); toast('Running... (⌘+Enter shortcut)','success'); }
      }
    }
  }
});
// Add ⌘+Enter hint to primary textareas
document.addEventListener('DOMContentLoaded',function(){
  document.querySelectorAll('.card .btn-primary').forEach(function(btn){
    var card=btn.closest('.card');
    if(card){
      var hint=document.createElement('p');
      hint.style.cssText='font-size:11px;color:#a8aab4;margin-top:8px;text-align:center;';
      hint.textContent='⌘+Enter to run';
      btn.parentNode.insertBefore(hint, btn.nextSibling);
    }
  });
});

/* ── Copy to clipboard ── */
function copyToClipboard(text){
  navigator.clipboard.writeText(text).then(function(){
    toast('Copied to clipboard!','success');
  }).catch(function(){ toast('Copy failed — select text manually','warn'); });
}

function renderOutput(result, dlKey, title, gmailSubject, isLinkedin){
  var gdoc='https://docs.google.com/document/create?title='+encodeURIComponent(title);
  var gmailBtn='', liBtn='';
  if(gmailSubject){
    var gu='https://mail.google.com/mail/?view=cm&su='+encodeURIComponent(gmailSubject)+'&body='+encodeURIComponent(result.slice(0,1500));
    gmailBtn='<a class="btn btn-secondary btn-sm" href="'+gu+'" target="_blank">✉️ Gmail</a>';
  }
  if(isLinkedin){
    var lu='https://www.linkedin.com/feed/?shareActive=true&text='+encodeURIComponent(result.slice(0,700));
    liBtn='<a class="btn btn-secondary btn-sm" href="'+lu+'" target="_blank">🔗 LinkedIn</a>';
  }
  return '<div class="output"><div class="output-label">AI Result</div>'
    +'<div class="output-text">'+result.replace(/\\n/g,'<br>')+'</div>'
    +'<div class="action-bar">'
    +'<a class="btn btn-secondary btn-sm" href="/download/'+dlKey+'">⬇ .docx</a>'
    +'<a class="btn btn-secondary btn-sm" href="/download-txt/'+dlKey+'">⬇ .txt</a>'
    +'<a class="btn btn-secondary btn-sm" href="'+gdoc+'" target="_blank">📄 Google Docs</a>'
    +gmailBtn+liBtn+'</div></div>';
}
function showAlert(targetId, msg, type){
  var cls = type==='error'?'alert-error':type==='warn'?'alert-warn':'alert-success';
  document.getElementById(targetId).innerHTML='<div class="alert '+cls+'">'+msg+'</div>';
}
</script>"""

# ── Tool HTML bodies ──────────────────────────────────────────────────────────
TOOL_BODIES = {}

TOOL_BODIES["resume"] = """
<div class="grid2">
  <div>
    <div class="card">
      <div class="card-title">Your Resume</div>
      <div class="upload-zone" onclick="document.getElementById('res-file').click()">
        <div class="upload-icon"><span class="material-symbols-outlined" style="font-size:40px;color:#0071e3;">upload_file</span></div>
        <p style="font-weight:600;margin-bottom:4px;">Drop your resume here</p>
        <p style="font-size:12px;color:#5e5e63;">PDF, DOCX up to 10MB</p>
        <input type="file" id="res-file" accept=".pdf,.docx" style="display:none" onchange="handleResFile(this)"/>
      </div>
      <p style="text-align:center;margin:12px 0;font-size:13px;color:#717785;">or paste below</p>
      <div class="form-group" style="margin:0;">
        <textarea id="res-text" rows="6" placeholder="Paste resume text here..."></textarea>
      </div>
    </div>
  </div>
  <div>
    <div class="card">
      <div class="card-title">Job Description</div>
      <div class="form-group">
        <textarea id="res-jd" rows="10" placeholder="Paste the job posting here..."></textarea>
      </div>
      <button class="btn btn-primary" onclick="runResume()">
        <span class="material-symbols-outlined">auto_fix_high</span> Analyze Resume
      </button>
    </div>
  </div>
</div>
<div id="res-output"></div>
<script>
function handleResFile(input){
  var file = input.files[0]; if(!file) return;
  var zone = input.closest('.upload-zone');
  if(zone) zone.innerHTML = '<div style="color:var(--primary-c);font-size:14px;font-weight:600;display:flex;align-items:center;gap:8px;"><span class=\'material-symbols-outlined\'>description</span>'+file.name+'</div><p style="font-size:12px;color:var(--secondary);margin-top:6px;">Parsing...</p>';
  var fd = new FormData(); fd.append('file', file);
  fetch('/api/parse-file',{method:'POST',body:fd})
  .then(r=>r.json()).then(d=>{
    if(d.error){ if(zone) zone.innerHTML='<p style="color:red">Error: '+d.error+'</p>'; return; }
    document.getElementById('res-text').value = d.text;
    if(zone) zone.innerHTML='<div style="color:#059669;font-size:14px;font-weight:600;display:flex;align-items:center;gap:8px;"><span class=\'material-symbols-outlined\'>check_circle</span>'+file.name+' — ready</div><p style="font-size:12px;color:var(--secondary);margin-top:4px;">'+d.text.length+' characters extracted</p>';
  }).catch(e=>{ if(zone) zone.innerHTML='<p style="color:red">Parse failed. Paste text below.</p>'; });
}
// Same for job finder file
function handleJfFile(input){
  var file = input.files[0]; if(!file) return;
  var zone = input.closest('.upload-zone');
  if(zone) zone.innerHTML = '<div style="color:var(--primary-c);font-size:14px;font-weight:600;">📄 '+file.name+' — Parsing...</div>';
  var fd = new FormData(); fd.append('file', file);
  fetch('/api/parse-file',{method:'POST',body:fd})
  .then(r=>r.json()).then(d=>{
    if(d.error){ return; }
    document.getElementById('jf-res').value = d.text;
    if(zone) zone.innerHTML='<div style="color:#059669;font-size:14px;font-weight:600;">✓ '+file.name+' ready</div>';
  });
}
function runResume(){
  var resume = document.getElementById('res-text').value;
  var jd = document.getElementById('res-jd').value;
  if(!resume || !jd){ showAlert('res-output','Please add your resume and job description.','warn'); return; }
  showSpinner();
  fetch('/api/resume',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({resume:resume,jd:jd,summary:'Resume reviewed'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){ showAlert('res-output',d.error,'error'); return; }
    document.getElementById('res-output').innerHTML = renderOutput(d.result,'resume','My Reviewed Resume','');
    toast('✓ Resume analyzed!','success');
  }).catch(e=>{ hideSpinner(); showAlert('res-output',e.message,'error'); });
}
</script>"""

TOOL_BODIES["interview"] = """
<div class="card" style="max-width:600px;">
  <div class="card-title">Setup Interview</div>
  <div class="grid3">
    <div class="form-group">
      <label>Job Role</label>
      <input type="text" id="int-role" placeholder="e.g. Data Analyst"/>
    </div>
    <div class="form-group">
      <label>Experience Level</label>
      <select id="int-level">
        <option>Entry (0-2 yrs)</option><option>Mid (2-5 yrs)</option><option>Senior (5+ yrs)</option>
      </select>
    </div>
    <div class="form-group">
      <label>Interview Type</label>
      <select id="int-type">
        <option>General / behavioral</option><option>Technical</option>
        <option>Case study</option><option>HR round</option>
      </select>
    </div>
  </div>
  <button class="btn btn-primary" onclick="startInterview()">
    <span class="material-symbols-outlined">play_arrow</span> Start Interview
  </button>
</div>
<div id="int-chat-area" style="display:none;max-width:700px;">
  <div class="card">
    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:16px;">
      <span style="font-weight:600;font-size:15px;">Live Interview</span>
      <button class="btn btn-secondary btn-sm" onclick="resetInterview()">Reset</button>
    </div>
    <div id="int-chat" class="chat-wrap"></div>
    <div style="display:flex;gap:8px;margin-top:12px;">
      <input type="text" id="int-ans" placeholder="Type your answer..." style="flex:1;"
             onkeydown="if(event.key==='Enter')sendAnswer()"/>
      <button class="btn btn-primary" style="width:auto;padding:10px 20px;" onclick="sendAnswer()">Send ↗</button>
    </div>
  </div>
</div>
<script>
var intHistory = [];
var intSystem = '';
function startInterview(){
  var role = document.getElementById('int-role').value;
  if(!role){ showAlert('int-chat-area','Enter a job role.','warn'); return; }
  var level = document.getElementById('int-level').value;
  var type  = document.getElementById('int-type').value;
  intSystem = role+'||'+level+'||'+type;
  intHistory = [];
  showSpinner();
  fetch('/api/interview_answer',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({answer:'Start the interview.',history:[],role:role,level:level,type:type,summary:'Interview started'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){ alert(d.error); return; }
    intHistory.push({role:'assistant',content:d.result});
    document.getElementById('int-chat-area').style.display='block';
    renderChat();
  });
}
function sendAnswer(){
  var ans = document.getElementById('int-ans').value.trim();
  if(!ans) return;
  intHistory.push({role:'user',content:ans});
  document.getElementById('int-ans').value='';
  renderChat();
  var parts = intSystem.split('||');
  showSpinner();
  fetch('/api/interview_answer',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({answer:ans,history:intHistory.slice(0,-1),role:parts[0],level:parts[1],type:parts[2],summary:'Interview answer'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){ alert(d.error); return; }
    intHistory.push({role:'assistant',content:d.result});
    renderChat();
  });
}
function renderChat(){
  var html='';var qCount=0;
  intHistory.forEach(function(m){
    if(m.role==='assistant'){
      qCount++;
      html+='<div class="bubble-ai"><span style="font-size:11px;font-weight:600;color:var(--secondary);letter-spacing:.05em;display:block;margin-bottom:6px;">QUESTION '+qCount+'</span>'+m.content.replace(/\\\\n/g,'<br>')+'</div>';
    } else if(m.content!=='Start the interview.'){
      html+='<div class="bubble-user">'+m.content+'</div>';
    }
  });
  document.getElementById('int-chat').innerHTML=html;
  var chat=document.getElementById('int-chat');
  chat.scrollTop=chat.scrollHeight;
});
  document.getElementById('int-chat').innerHTML=html;
  document.getElementById('int-chat').scrollTop=99999;
}
function resetInterview(){ intHistory=[]; document.getElementById('int-chat').innerHTML=''; document.getElementById('int-chat-area').style.display='none'; }
</script>"""

TOOL_BODIES["cover"] = """
<div class="grid2">
  <div>
    <div class="card">
      <div class="card-title">Job Details</div>
      <div class="form-group"><label>Job Description</label><textarea id="cov-jd" rows="7" placeholder="Paste the job posting..."></textarea></div>
      <div class="form-group"><label>Company Name (optional)</label><input type="text" id="cov-co" placeholder="e.g. Google, Infosys"/></div>
      <div class="form-group"><label>Tone</label>
        <select id="cov-tone">
          <option>Professional & formal</option><option>Friendly & conversational</option>
          <option>Confident & bold</option><option>Humble & enthusiastic</option>
        </select>
      </div>
    </div>
  </div>
  <div>
    <div class="card">
      <div class="card-title">About You</div>
      <div class="form-group"><label>Full Name</label><input type="text" id="cov-name" placeholder="Your full name"/></div>
      <div class="form-group"><label>Email</label><input type="text" id="cov-email" placeholder="your@email.com"/></div>
      <div class="form-group"><label>Skills & Experience</label><textarea id="cov-about" rows="5" placeholder="Your key skills, achievements, experience..."></textarea></div>
      <button class="btn btn-primary" onclick="runCover()">
        <span class="material-symbols-outlined">auto_fix_high</span> Generate Cover Letter
      </button>
    </div>
  </div>
</div>
<div id="cov-output"></div>
<script>
function runCover(){
  var jd=document.getElementById('cov-jd').value, about=document.getElementById('cov-about').value;
  if(!jd||!about){showAlert('cov-output','Fill in job description and your details.','warn');return;}
  showSpinner();
  fetch('/api/cover',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({jd:jd,about:about,name:document.getElementById('cov-name').value,
      email:document.getElementById('cov-email').value,company:document.getElementById('cov-co').value,
      tone:document.getElementById('cov-tone').value,summary:'Cover letter generated'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){showAlert('cov-output',d.error,'error');return;}
    var co=document.getElementById('cov-co').value||'Job';
    document.getElementById('cov-output').innerHTML=renderOutput(d.result,'cover','Cover Letter','Application – '+co);
    toast('✓ Cover letter ready!','success');
  });
}
</script>"""

TOOL_BODIES["linkedin"] = """
<div class="card" style="max-width:680px;">
  <div class="card-title">Create LinkedIn Post</div>
  <div class="grid2">
    <div class="form-group"><label>Post Type</label>
      <select id="li-type">
        <option>Career milestone</option><option>New job announcement</option>
        <option>Lesson learned</option><option>Project showcase</option>
        <option>Thought leadership</option><option>Gratitude post</option>
      </select>
    </div>
    <div class="form-group"><label>Writing Style</label>
      <select id="li-style">
        <option>Storytelling</option><option>Professional</option>
        <option>Casual & relatable</option><option>Inspirational</option>
      </select>
    </div>
  </div>
  <div class="form-group">
    <label>What's it about?</label>
    <textarea id="li-topic" rows="5" placeholder="Describe your achievement, project, or thought in your own words..."></textarea>
  </div>
  <button class="btn btn-primary" onclick="runLinkedin()">
    <span class="material-symbols-outlined">auto_fix_high</span> Generate Post
  </button>
</div>
<div id="li-output"></div>
<script>
function runLinkedin(){
  var topic=document.getElementById('li-topic').value;
  if(!topic){showAlert('li-output','Describe what you want to post about.','warn');return;}
  showSpinner();
  fetch('/api/linkedin',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({topic:topic,ptype:document.getElementById('li-type').value,
      style:document.getElementById('li-style').value,summary:'LinkedIn post'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){showAlert('li-output',d.error,'error');return;}
    document.getElementById('li-output').innerHTML=renderOutput(d.result,'linkedin','LinkedIn Post','',true);
    toast('✓ LinkedIn post ready!','success');
  });
}
</script>"""

TOOL_BODIES["decoder"] = """
<div class="card" style="max-width:780px;">
  <div class="card-title">Paste Job Posting</div>
  <div class="form-group">
    <textarea id="dec-jd" rows="10" placeholder="Paste the full job description here..."></textarea>
  </div>
  <button class="btn btn-primary" onclick="runDecoder()">
    <span class="material-symbols-outlined">analytics</span> Decode This Job
  </button>
</div>
<div id="dec-output"></div>
<script>
function runDecoder(){
  var jd=document.getElementById('dec-jd').value;
  if(!jd){showAlert('dec-output','Paste a job description first.','warn');return;}
  showSpinner();
  fetch('/api/decoder',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({jd:jd,summary:'Job decoded'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){showAlert('dec-output',d.error,'error');return;}
    document.getElementById('dec-output').innerHTML=renderOutput(d.result,'decoder','Job Analysis','');
    toast('✓ Job decoded!','success');
  });
}
</script>"""

TOOL_BODIES["apply"] = """
<div class="grid2">
  <div>
    <div class="card">
      <div class="card-title">Job Details</div>
      <div class="form-group"><label>Job Description</label><textarea id="ap-jd" rows="8" placeholder="Paste job posting..."></textarea></div>
      <div class="form-group"><label>Job URL (optional)</label><input type="text" id="ap-url" placeholder="https://..."/></div>
      <div class="form-group"><label>HR Email (optional)</label><input type="text" id="ap-hr" placeholder="hr@company.com"/></div>
    </div>
  </div>
  <div>
    <div class="card">
      <div class="card-title">Your Details</div>
      <div class="form-group"><label>Your Resume / Key Info</label><textarea id="ap-res" rows="6" placeholder="Paste resume or key experience..."></textarea></div>
      <div class="form-group"><label>Full Name</label><input type="text" id="ap-name" placeholder="Your name"/></div>
      <div class="form-group"><label>Email</label><input type="text" id="ap-email" placeholder="your@email.com"/></div>
      <button class="btn btn-primary" onclick="runApply()">
        <span class="material-symbols-outlined">rocket_launch</span> Generate Package
      </button>
    </div>
  </div>
</div>
<div id="ap-output"></div>
<script>
function runApply(){
  var jd=document.getElementById('ap-jd').value, res=document.getElementById('ap-res').value;
  if(!jd||!res){showAlert('ap-output','Fill in job description and your details.','warn');return;}
  showSpinner();
  fetch('/api/apply',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({jd:jd,resume:res,name:document.getElementById('ap-name').value,
      email:document.getElementById('ap-email').value,url:document.getElementById('ap-url').value,
      hr:document.getElementById('ap-hr').value,summary:'Application package'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){showAlert('ap-output',d.error,'error');return;}
    var pkg=JSON.parse(d.result);
    var jdQ=encodeURIComponent(document.getElementById('ap-jd').value.slice(0,50));
    var hrEmail=document.getElementById('ap-hr').value;
    var name=document.getElementById('ap-name').value;
    var appUrl=document.getElementById('ap-url').value;
    var gmailLink=hrEmail?'<a class="apply-link" href="https://mail.google.com/mail/?view=cm&to='+hrEmail+'&su='+encodeURIComponent('Job Application – '+name)+'&body='+encodeURIComponent(pkg.email.slice(0,1000))+'" target="_blank">✉️ Gmail</a>':'';
    var siteLink=appUrl?'<a class="apply-link" href="'+appUrl+'" target="_blank">🔗 Company Site</a>':'';
    toast('✓ Application package ready!','success');
    document.getElementById('ap-output').innerHTML=`
      <div class="card"><div class="card-title">✉️ Cover Letter</div><div class="output-text">${pkg.cover.replace(/\\n/g,'<br>')}</div></div>
      <div class="card"><div class="card-title">📧 Application Email</div><div class="output-text">${pkg.email.replace(/\\n/g,'<br>')}</div></div>
      <div class="card"><div class="card-title">💡 Key Talking Points</div><div class="output-text">${pkg.points.replace(/\\n/g,'<br>')}</div></div>
      <div class="apply-box"><h4 style="font-size:17px;font-weight:600;margin-bottom:12px;">Quick Apply</h4>
        ${siteLink}${gmailLink}
        <a class="apply-link" href="https://www.linkedin.com/jobs/search/?keywords=${jdQ}" target="_blank">💼 LinkedIn</a>
        <a class="apply-link" href="https://www.naukri.com" target="_blank">🇮🇳 Naukri</a>
        <a class="apply-link" href="https://www.shine.com" target="_blank">✨ Shine</a>
      </div>`;
  });
}
</script>"""

TOOL_BODIES["jobs"] = """
<div class="grid2">
  <div>
    <div class="card">
      <div class="card-title">Your Resume</div>
      <div class="upload-zone" onclick="document.getElementById('jf-file').click()" style="padding:20px;margin-bottom:12px;">
        <span class="material-symbols-outlined" style="font-size:32px;color:#0071e3;display:block;margin-bottom:8px;">upload_file</span>
        <p style="font-weight:600;font-size:14px;">Upload Resume (PDF/DOCX)</p>
        <p style="font-size:12px;color:#5e5e63;margin-top:4px;">or paste below</p>
        <input type="file" id="jf-file" accept=".pdf,.docx" style="display:none" onchange="handleJfFile(this)"/>
      </div>
      <div class="form-group"><textarea id="jf-res" rows="5" placeholder="Paste your resume text here..."></textarea></div>
    </div>
  </div>
  <div>
    <div class="card">
      <div class="card-title">Preferences</div>
      <div class="form-group"><label>Job Role</label><input type="text" id="jf-role" placeholder="e.g. Software Engineer, Product Manager"/></div>
      <div class="form-group"><label>Location</label><input type="text" id="jf-loc" placeholder="e.g. Mumbai, Bangalore, Remote"/></div>
      <div class="form-group"><label>Experience</label>
        <select id="jf-exp">
          <option>Fresher (0-1 yr)</option><option>Junior (1-3 yrs)</option>
          <option>Mid (3-6 yrs)</option><option>Senior (6+ yrs)</option>
        </select>
      </div>
      <div class="form-group"><label>Job Type</label>
        <select id="jf-type">
          <option>Full-time</option><option>Part-time</option>
          <option>Remote</option><option>Freelance</option><option>Internship</option>
        </select>
      </div>
      <button class="btn btn-primary" onclick="runJobs()">
        <span class="material-symbols-outlined">search</span> Find Matching Jobs
      </button>
    </div>
  </div>
</div>
<div id="jf-output"></div>
<script>
function runJobs(){
  var res=document.getElementById('jf-res').value, role=document.getElementById('jf-role').value;
  if(!res&&!role){showAlert('jf-output','Add your resume or job role.','warn');return;}
  showSpinner();
  fetch('/api/jobs',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({resume:res,role:role,location:document.getElementById('jf-loc').value,
      exp:document.getElementById('jf-exp').value,summary:'Job search'})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){showAlert('jf-output',d.error,'error');return;}
    var roleQ=encodeURIComponent(role||'software engineer');
    var locQ=encodeURIComponent(document.getElementById('jf-loc').value||'India');
    var roleSlug=(role||'software-engineer').toLowerCase().replace(/ /g,'-');
    var locSlug=(document.getElementById('jf-loc').value||'india').toLowerCase().replace(/ /g,'-');
    toast('✓ Jobs found!','success');
    document.getElementById('jf-output').innerHTML=`
      <div class="output"><div class="output-label">Profile Analysis</div>
        <div class="output-text">${d.result.replace(/\\n/g,'<br>')}</div></div>
      <div class="apply-box" style="margin-top:16px;">
        <h4 style="font-size:17px;font-weight:600;margin-bottom:12px;">Search on Job Portals</h4>
        <a class="apply-link" href="https://www.naukri.com/${roleSlug}-jobs-in-${locSlug}" target="_blank">🇮🇳 Naukri</a>
        <a class="apply-link" href="https://www.linkedin.com/jobs/search/?keywords=${roleQ}&location=${locQ}" target="_blank">💼 LinkedIn</a>
        <a class="apply-link" href="https://www.indeed.com/jobs?q=${roleQ}&l=${locQ}" target="_blank">🔍 Indeed</a>
        <a class="apply-link" href="https://www.shine.com/job-search/${roleSlug}-jobs" target="_blank">✨ Shine</a>
        <a class="apply-link" href="https://internshala.com/jobs/${roleSlug}-jobs" target="_blank">🎓 Internshala</a>
        <a class="apply-link" href="https://www.glassdoor.co.in/Job/jobs.htm?sc.keyword=${roleQ}" target="_blank">🏢 Glassdoor</a>
        <a class="apply-link" href="https://wellfound.com/jobs?query=${roleQ}" target="_blank">🚀 Wellfound</a>
      </div>`;
  });
}
</script>"""

TOOL_BODIES["career"] = """
<div class="grid2">
  <div>
    <div class="card">
      <div class="card-title">Where You Are Now</div>
      <div class="form-group"><label>Current Job Role</label><input type="text" id="cp-role" placeholder="e.g. Junior Software Engineer"/></div>
      <div class="form-group"><label>Years of Experience</label><input type="text" id="cp-exp" placeholder="e.g. 2 years"/></div>
      <div class="form-group"><label>Your Current Skills</label><textarea id="cp-skills" rows="4" placeholder="e.g. Python, React, SQL, project management..."></textarea></div>
    </div>
  </div>
  <div>
    <div class="card">
      <div class="card-title">Where You Want to Go</div>
      <div class="form-group"><label>Dream Job / Goal (optional)</label><input type="text" id="cp-goal" placeholder="e.g. CTO, Product Manager, Data Scientist"/></div>
      <div class="form-group"><label>Industry</label>
        <select id="cp-ind">
          <option>Technology</option><option>Finance</option><option>Marketing</option>
          <option>Healthcare</option><option>Education</option><option>E-commerce</option>
          <option>Consulting</option><option>Government</option><option>Other</option>
        </select>
      </div>
      <div class="form-group"><label>Your Timeline</label>
        <select id="cp-time">
          <option>6 months</option><option>1 year</option><option>2 years</option>
          <option>3-5 years</option><option>Long term (5+ years)</option>
        </select>
      </div>
      <button class="btn btn-primary" onclick="runCareer()">
        <span class="material-symbols-outlined">map</span> Show My Career Path
      </button>
    </div>
  </div>
</div>
<div id="cp-output"></div>
<script>
function runCareer(){
  var role=document.getElementById('cp-role').value;
  if(!role){showAlert('cp-output','Enter your current job role.','warn');return;}
  showSpinner();
  fetch('/api/career',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({role:role,exp:document.getElementById('cp-exp').value,
      skills:document.getElementById('cp-skills').value,goal:document.getElementById('cp-goal').value,
      industry:document.getElementById('cp-ind').value,timeline:document.getElementById('cp-time').value,
      summary:'Career path for '+role})})
  .then(r=>r.json()).then(d=>{
    hideSpinner();
    if(d.error){showAlert('cp-output',d.error,'error');return;}
    document.getElementById('cp-output').innerHTML=renderOutput(d.result,'career','Career Path Roadmap','');
    toast('✓ Career roadmap ready!','success');
  });
}
</script>"""

TOOL_BODIES["salary"] = """
<div class="tabs">
  <button class="tab-btn active" data-group="sal" onclick="switchTab('sal',0)">💰 Know Your Worth</button>
  <button class="tab-btn" data-group="sal" onclick="switchTab('sal',1)">🗣 Negotiation Script</button>
  <button class="tab-btn" data-group="sal" onclick="switchTab('sal',2)">📧 Counter Offer Email</button>
</div>

<div class="tab-content active" data-group="sal">
  <div class="grid2">
    <div class="card">
      <div class="form-group"><label>Job Role</label><input type="text" id="sw-role" placeholder="e.g. Senior Data Analyst"/></div>
      <div class="form-group"><label>Years of Experience</label><input type="text" id="sw-exp" placeholder="e.g. 4 years"/></div>
      <div class="form-group"><label>City</label><input type="text" id="sw-city" placeholder="e.g. Bangalore, Mumbai, Delhi"/></div>
    </div>
    <div class="card">
      <div class="form-group"><label>Your Key Skills</label><textarea id="sw-skills" rows="3" placeholder="e.g. Python, ML, SQL..."></textarea></div>
      <div class="form-group"><label>Company (optional)</label><input type="text" id="sw-co" placeholder="e.g. Infosys, Google, Swiggy"/></div>
      <button class="btn btn-primary" onclick="runSalaryWorth()">
        <span class="material-symbols-outlined">payments</span> Check My Market Value
      </button>
    </div>
  </div>
  <div id="sw-output"></div>
</div>

<div class="tab-content" data-group="sal">
  <div class="grid2">
    <div class="card">
      <div class="form-group"><label>Offer Received (LPA)</label><input type="text" id="ns-offer" placeholder="e.g. 8 LPA"/></div>
      <div class="form-group"><label>What You Want (LPA)</label><input type="text" id="ns-want" placeholder="e.g. 12 LPA"/></div>
      <div class="form-group"><label>Job Role</label><input type="text" id="ns-role" placeholder="e.g. Product Manager"/></div>
    </div>
    <div class="card">
      <div class="form-group"><label>Experience</label><input type="text" id="ns-exp" placeholder="e.g. 5 years"/></div>
      <div class="form-group"><label>Your Achievements</label><textarea id="ns-ach" rows="4" placeholder="e.g. Led team of 10, increased revenue 40%..."></textarea></div>
      <button class="btn btn-primary" onclick="runScript()">
        <span class="material-symbols-outlined">record_voice_over</span> Write My Script
      </button>
    </div>
  </div>
  <div id="ns-output"></div>
</div>

<div class="tab-content" data-group="sal">
  <div class="grid2">
    <div class="card">
      <div class="form-group"><label>Your Name</label><input type="text" id="ce-name" placeholder="Your full name"/></div>
      <div class="form-group"><label>Offer Received (LPA)</label><input type="text" id="ce-offer" placeholder="e.g. 8 LPA"/></div>
      <div class="form-group"><label>Counter Offer (LPA)</label><input type="text" id="ce-want" placeholder="e.g. 12 LPA"/></div>
    </div>
    <div class="card">
      <div class="form-group"><label>Role</label><input type="text" id="ce-role" placeholder="e.g. Senior Engineer"/></div>
      <div class="form-group"><label>Company</label><input type="text" id="ce-co" placeholder="e.g. Infosys"/></div>
      <div class="form-group"><label>Why You Deserve More</label><textarea id="ce-why" rows="3" placeholder="5 yrs experience, led major projects..."></textarea></div>
      <button class="btn btn-primary" onclick="runCounterEmail()">
        <span class="material-symbols-outlined">mail</span> Write Counter Offer Email
      </button>
    </div>
  </div>
  <div id="ce-output"></div>
</div>

<script>
function runSalaryWorth(){
  var role=document.getElementById('sw-role').value;
  if(!role){showAlert('sw-output','Enter a job role.','warn');return;}
  showSpinner();
  fetch('/api/salary_worth',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({role:role,exp:document.getElementById('sw-exp').value,
      city:document.getElementById('sw-city').value,skills:document.getElementById('sw-skills').value,
      company:document.getElementById('sw-co').value,summary:'Salary worth for '+role})})
  .then(r=>r.json()).then(d=>{hideSpinner();
    if(d.error){showAlert('sw-output',d.error,'error');return;}
    document.getElementById('sw-output').innerHTML=renderOutput(d.result,'salary_worth','Salary Analysis','');
    toast('✓ Salary analysis ready!','success');});
}
function runScript(){
  var offer=document.getElementById('ns-offer').value,want=document.getElementById('ns-want').value;
  if(!offer||!want){showAlert('ns-output','Enter offer and expected salary.','warn');return;}
  showSpinner();
  fetch('/api/salary_script',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({offer:offer,expect:want,role:document.getElementById('ns-role').value,
      exp:document.getElementById('ns-exp').value,achievements:document.getElementById('ns-ach').value,
      summary:'Negotiation script'})})
  .then(r=>r.json()).then(d=>{hideSpinner();
    if(d.error){showAlert('ns-output',d.error,'error');return;}
    document.getElementById('ns-output').innerHTML=renderOutput(d.result,'salary_script','Negotiation Script','');});
}
function runCounterEmail(){
  var offer=document.getElementById('ce-offer').value,want=document.getElementById('ce-want').value;
  if(!offer||!want){showAlert('ce-output','Enter offer and counter offer amounts.','warn');return;}
  var co=document.getElementById('ce-co').value, role=document.getElementById('ce-role').value;
  showSpinner();
  fetch('/api/salary_email',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({name:document.getElementById('ce-name').value,offer:offer,counter:want,
      role:role,company:co,reason:document.getElementById('ce-why').value,
      summary:'Counter offer email'})})
  .then(r=>r.json()).then(d=>{hideSpinner();
    if(d.error){showAlert('ce-output',d.error,'error');return;}
    document.getElementById('ce-output').innerHTML=renderOutput(d.result,'salary_email','Counter Offer Email',
      'Counter Offer – '+role+' – '+co);});
}
</script>"""

TOOL_BODIES["history"] = """{% set hist = session.get('history', []) %}
{% if not hist %}
<div class="empty">
  <div class="empty-icon"><span class="material-symbols-outlined" style="font-size:52px;color:#c1c6d6;">history</span></div>
  <div class="empty-title">No history yet</div>
  <p>Use any tool and your results will be saved here automatically.</p>
</div>
{% else %}
<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:20px;">
  <span style="font-weight:600;">{{ hist|length }} saved items</span>
  <form method="POST" action="/history/clear">
    <button type="submit" class="btn btn-danger btn-sm">🗑 Clear All</button>
  </form>
</div>
{% for entry in hist %}
<div class="history-item">
  <div class="history-head" onclick="this.nextElementSibling.classList.toggle('open')">
    <span>{{ entry.tool }} &nbsp;·&nbsp; {{ entry.time }} — {{ entry.summary[:80] }}...</span>
    <span class="material-symbols-outlined" style="font-size:18px;color:#717785;">expand_more</span>
  </div>
  <div class="history-body {% if loop.first %}open{% endif %}">
    <div class="output-text" style="margin-bottom:16px;">{{ entry.output | replace('\n','<br>') | safe }}</div>
    <div class="action-bar">
      <a class="btn btn-secondary btn-sm" href="/download/hist_{{ loop.index0 }}">⬇ .docx</a>
      <a class="btn btn-secondary btn-sm" href="/download-txt/hist_{{ loop.index0 }}">⬇ .txt</a>
      <a class="btn btn-secondary btn-sm" href="https://docs.google.com/document/create?title={{ entry.tool | urlencode }}" target="_blank">📄 Google Docs</a>
    </div>
  </div>
</div>
{% endfor %}
{% endif %}"""

# Fix history download to use session history
@app.route("/download/hist_<int:idx>")
def download_hist_docx(idx):
    hist = session.get("history", [])
    if idx >= len(hist): return "Not found", 404
    entry = hist[idx]
    buf = mk_docx(entry["tool"], entry["output"])
    return send_file(buf, as_attachment=True, download_name=f"history_{idx}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/download-txt/hist_<int:idx>")
def download_hist_txt(idx):
    hist = session.get("history", [])
    if idx >= len(hist): return "Not found", 404
    return send_file(io.BytesIO(hist[idx]["output"].encode()), as_attachment=True,
                     download_name=f"history_{idx}.txt", mimetype="text/plain")

# Inject shared JS before </body>
# BASE_WITH_JS built after SHARED_JS is defined below

@app.errorhandler(404)
def not_found(e):
    from flask import redirect
    return redirect("/")

@app.errorhandler(500)
def server_error(e):
    return f'<div style="padding:2rem;font-family:sans-serif;"><h2>Error</h2><p>{str(e)}</p><a href="/">← Back</a></div>', 500

if __name__ == "__main__":
    print("\n🚀 CareerAI Flask App")
    print("   Open: http://127.0.0.1:5000\n")
    app.run(debug=True, host="127.0.0.1", port=5000)
