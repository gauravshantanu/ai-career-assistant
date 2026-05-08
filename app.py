import streamlit as st
from groq import Groq
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
import urllib.parse

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="AI Career Assistant", page_icon="💼", layout="wide", initial_sidebar_state="collapsed")

# ── CSS ── Stitch Pro Digital Interface ───────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Material+Symbols+Outlined:opsz,wght,FILL,GRAD@24,400,0,0&display=swap');

:root {
  --surface:#f9f9ff;
  --surface-dim:#d8d9e3;
  --surface-container-lowest:#ffffff;
  --surface-container-low:#f2f3fd;
  --surface-container:#ecedf7;
  --surface-container-high:#e6e8f2;
  --surface-container-highest:#e0e2ec;
  --on-surface:#181c23;
  --on-surface-variant:#414753;
  --outline:#717785;
  --outline-variant:#c1c6d6;
  --primary:#0059b5;
  --on-primary:#ffffff;
  --primary-container:#0071e3;
  --on-primary-container:#fcfbff;
  --secondary:#5e5e63;
  --on-secondary:#ffffff;
  --secondary-container:#e0dfe4;
  --on-secondary-container:#626267;
  --tertiary:#9b3f00;
  --tertiary-container:#c25100;
  --on-tertiary-container:#fffaf9;
  --tertiary-fixed:#ffdbcb;
  --error:#ba1a1a;
  --background:#f5f5f7;
  --on-background:#181c23;
  --nav-height:52px;
  --sidebar-width:256px;
  --font:Inter,-apple-system,BlinkMacSystemFont,"Helvetica Neue",sans-serif;
}

/* ── Reset & Base ── */
*{box-sizing:border-box;margin:0;padding:0;}
html,body,[class*="css"]{
  font-family:var(--font)!important;
  background:var(--background)!important;
  color:var(--on-background)!important;
  -webkit-font-smoothing:antialiased!important;
}
.main,[data-testid="stAppViewContainer"],[data-testid="stMain"]{background:var(--background)!important;}
.block-container{
  padding:calc(var(--nav-height) + 32px) 2rem 4rem calc(var(--sidebar-width) + 2rem)!important;
  max-width:100%!important;
}
#MainMenu,footer,[data-testid="stToolbar"]{visibility:hidden!important;height:0!important;}
[data-testid="stDecoration"],[data-testid="collapsedControl"]{display:none!important;}
section[data-testid="stSidebar"]{display:none!important;}

/* ── Glass Nav Bar ── */
.stitch-nav{
  position:fixed;top:0;left:0;right:0;z-index:9999;
  height:var(--nav-height);
  background:rgba(249,249,255,0.8);
  backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);
  border-bottom:0.5px solid var(--outline-variant);
  display:flex;align-items:center;
  justify-content:space-between;
  padding:0 24px;
}
.stitch-nav-brand{
  font-size:24px;font-weight:600;color:var(--on-surface);
  letter-spacing:-0.011em;line-height:1.3;
  font-family:var(--font);white-space:nowrap;
}
.stitch-nav-links{display:flex;align-items:center;gap:24px;margin-left:32px;}
.stitch-nav-link{
  color:var(--secondary);font-size:15px;font-weight:500;
  text-decoration:none;letter-spacing:-0.015em;
  transition:color .2s;
}
.stitch-nav-link:hover{color:var(--primary);}
.stitch-nav-link.on{color:var(--primary);font-weight:600;border-bottom:2px solid var(--primary);padding-bottom:2px;}
.stitch-nav-right{display:flex;align-items:center;gap:12px;margin-left:auto;}
.stitch-upgrade-btn{
  background:var(--primary-container);color:var(--on-primary-container);
  padding:4px 16px;border-radius:8px;font-size:13px;font-weight:600;
  border:none;cursor:pointer;letter-spacing:0.01em;
  transition:opacity .15s;
}
.stitch-upgrade-btn:hover{opacity:.9;}
.stitch-icon-btn{
  background:transparent;border:none;cursor:pointer;
  color:var(--on-surface-variant);font-size:22px;
  display:flex;align-items:center;transition:color .15s;
}
.stitch-icon-btn:hover{color:var(--primary);}
.stitch-avatar{
  width:32px;height:32px;border-radius:50%;
  background:var(--surface-container);
  border:1px solid var(--outline-variant);
  display:flex;align-items:center;justify-content:center;
  font-size:12px;font-weight:600;color:var(--primary);
}

/* ── Left Sidebar ── */
.stitch-sidebar{
  position:fixed;left:0;top:var(--nav-height);bottom:0;
  width:var(--sidebar-width);
  background:var(--surface-container-low);
  border-right:0.5px solid var(--outline-variant);
  overflow-y:auto;padding:16px;z-index:100;
  display:flex;flex-direction:column;gap:4px;
}
.stitch-sidebar-header{
  margin-bottom:16px;padding:0 4px;
}
.stitch-sidebar-title{
  font-size:20px;font-weight:600;color:var(--primary);
  letter-spacing:-0.007em;line-height:1.4;
}
.stitch-sidebar-sub{
  font-size:13px;font-weight:500;color:var(--secondary);
  letter-spacing:0.01em;margin-top:2px;
}
.stitch-nav-item{
  display:flex;align-items:center;gap:8px;
  padding:8px 12px;border-radius:8px;
  text-decoration:none;color:var(--on-surface-variant);
  font-size:13px;font-weight:500;letter-spacing:0.01em;
  transition:all .15s;cursor:pointer;
}
.stitch-nav-item:hover{background:var(--surface-container-highest);color:var(--on-surface);}
.stitch-nav-item.active{
  background:var(--secondary-container);
  color:var(--on-secondary-container);font-weight:700;
}
.stitch-nav-item .ms{font-size:20px;}

/* ── Page header ── */
.stitch-page-header{margin-bottom:32px;}
.stitch-page-title{
  font-size:48px;font-weight:700;color:var(--on-surface);
  letter-spacing:-0.022em;line-height:1.1;margin-bottom:8px;
}
.stitch-page-sub{
  font-size:17px;font-weight:400;color:var(--secondary);
  letter-spacing:-0.022em;line-height:1.5;
}

/* ── Cards ── */
.stitch-card{
  background:var(--surface-container-lowest);
  border-radius:16px;
  border:1px solid rgba(193,198,214,0.2);
  box-shadow:0 4px 24px rgba(0,0,0,0.04);
  padding:24px;margin-bottom:16px;
}
.stitch-card-title{
  font-size:20px;font-weight:600;color:var(--on-surface);
  letter-spacing:-0.007em;line-height:1.4;margin-bottom:16px;
}

/* ── Output card ── */
.output-section{
  background:var(--surface-container-lowest);
  border-radius:16px;
  border:1px solid rgba(193,198,214,0.2);
  box-shadow:0 4px 24px rgba(0,0,0,0.04);
  padding:24px;margin-top:16px;
  font-size:15px;line-height:1.5;
  letter-spacing:-0.015em;
  white-space:pre-wrap;color:var(--on-surface);
}

/* ── Inputs ── */
.stTextArea textarea{
  background:var(--surface-container-lowest)!important;
  border:1px solid var(--outline-variant)!important;
  border-radius:8px!important;
  color:var(--on-surface)!important;
  font-family:var(--font)!important;
  font-size:15px!important;letter-spacing:-0.015em!important;
  line-height:1.5!important;
  transition:border .15s,box-shadow .15s!important;
}
.stTextArea textarea:focus{
  border-color:var(--primary-container)!important;
  box-shadow:0 0 0 2px rgba(0,113,227,0.15)!important;
}
.stTextInput input{
  background:var(--surface-container-lowest)!important;
  border:1px solid var(--outline-variant)!important;
  border-radius:8px!important;color:var(--on-surface)!important;
  font-family:var(--font)!important;font-size:15px!important;
  letter-spacing:-0.015em!important;
  transition:border .15s,box-shadow .15s!important;
}
.stTextInput input:focus{
  border-color:var(--primary-container)!important;
  box-shadow:0 0 0 2px rgba(0,113,227,0.15)!important;
}
.stSelectbox>div>div{
  background:var(--surface-container-lowest)!important;
  border:1px solid var(--outline-variant)!important;
  border-radius:8px!important;
  color:var(--on-surface)!important;
  font-family:var(--font)!important;font-size:15px!important;
}
label{
  color:var(--secondary)!important;font-family:var(--font)!important;
  font-size:13px!important;font-weight:500!important;letter-spacing:0.01em!important;
}

/* ── Buttons ── */
.stButton>button{
  background:var(--primary-container)!important;
  color:var(--on-primary-container)!important;
  border:none!important;border-radius:8px!important;
  font-family:var(--font)!important;font-weight:600!important;
  font-size:15px!important;padding:10px 20px!important;
  letter-spacing:-0.015em!important;
  transition:opacity .15s,transform .1s!important;
  box-shadow:0 4px 24px rgba(0,0,0,0.04)!important;
}
.stButton>button:hover{opacity:.9!important;}
.stButton>button:active{transform:scale(0.98)!important;}
.stDownloadButton>button{
  background:var(--secondary-container)!important;
  color:var(--on-secondary-container)!important;
  border:none!important;border-radius:8px!important;
  font-size:13px!important;font-weight:500!important;
  padding:6px 16px!important;letter-spacing:0.01em!important;
  transition:all .15s!important;
}
.stDownloadButton>button:hover{background:var(--surface-container-highest)!important;}

/* ── Tabs ── */
[data-testid="stTabs"] [data-baseweb="tab-list"]{
  background:var(--surface-container)!important;
  border-radius:8px!important;padding:4px!important;
  gap:2px!important;border:none!important;
}
[data-testid="stTabs"] [data-baseweb="tab"]{
  background:transparent!important;border:none!important;
  border-radius:6px!important;color:var(--secondary)!important;
  font-family:var(--font)!important;font-size:13px!important;
  font-weight:500!important;letter-spacing:0.01em!important;
  padding:6px 16px!important;
}
[data-testid="stTabs"] [aria-selected="true"]{
  background:var(--primary-container)!important;
  color:var(--on-primary-container)!important;font-weight:600!important;
  box-shadow:0 2px 8px rgba(0,0,0,0.08)!important;
}

/* ── Chat bubbles ── */
.chat-ai{
  background:var(--surface-container-lowest);
  border:0.5px solid var(--outline-variant);
  border-radius:12px 12px 12px 4px;
  box-shadow:0 4px 24px rgba(0,0,0,0.04);
  color:var(--on-surface);padding:14px 18px;
  margin:8px 15% 8px 0;font-size:15px;line-height:1.5;
  letter-spacing:-0.015em;
}
.chat-user{
  background:var(--primary-container);color:var(--on-primary-container);
  border-radius:12px 12px 4px 12px;
  padding:14px 18px;margin:8px 0 8px 15%;
  font-size:15px;line-height:1.5;letter-spacing:-0.015em;
  font-weight:500;
}

/* ── Apply/links box ── */
.apply-box{
  background:var(--surface-container-lowest);
  border-radius:16px;border:1px solid rgba(193,198,214,0.2);
  box-shadow:0 4px 24px rgba(0,0,0,0.04);
  padding:24px;margin-top:16px;
}
.apply-box h4{
  color:var(--on-surface)!important;font-family:var(--font)!important;
  font-size:20px!important;font-weight:600!important;
  letter-spacing:-0.007em!important;margin-bottom:12px!important;
}
.apply-link{
  display:inline-flex;align-items:center;gap:6px;
  background:var(--secondary-container);
  color:var(--on-secondary-container);
  border:none;padding:8px 16px;border-radius:8px;
  font-size:13px;font-weight:500;text-decoration:none;
  margin:4px 3px;letter-spacing:0.01em;
  transition:all .15s;
}
.apply-link:hover{background:var(--surface-container-highest);color:var(--on-surface);}

/* ── Language pills ── */
.lang-pill{
  display:inline-block;background:var(--secondary-container);
  color:var(--on-secondary-container);
  padding:4px 12px;border-radius:9999px;
  font-size:12px;font-weight:500;margin:3px;letter-spacing:0.01em;
}

/* ── Settings expander ── */
[data-testid="stExpander"]{
  background:var(--surface-container-lowest)!important;
  border:0.5px solid var(--outline-variant)!important;
  border-radius:8px!important;margin-bottom:16px!important;
}
.streamlit-expanderHeader{
  background:var(--surface-container-lowest)!important;
  color:var(--on-surface)!important;border-radius:8px!important;
  font-family:var(--font)!important;font-size:13px!important;
  font-weight:500!important;
}

/* ── Status messages ── */
.stSuccess{background:rgba(0,89,181,0.06)!important;border-radius:8px!important;}
.stError{background:rgba(186,26,26,0.06)!important;border-radius:8px!important;}
.stWarning{background:rgba(155,63,0,0.06)!important;border-radius:8px!important;}
.stSpinner>div{border-top-color:var(--primary-container)!important;}
hr{border-color:var(--outline-variant)!important;}

/* ── File uploader ── */
[data-testid="stFileUploadDropzone"]{
  background:var(--surface-container-low)!important;
  border:2px dashed var(--outline-variant)!important;
  border-radius:12px!important;
  transition:border-color .15s!important;
}
[data-testid="stFileUploadDropzone"]:hover{border-color:var(--primary-container)!important;}

/* ── Tool header ── */
.tool-header{display:flex;align-items:center;gap:12px;margin-bottom:4px;}
.tool-icon{
  width:40px;height:40px;background:rgba(0,113,227,0.1);
  border-radius:10px;display:flex;align-items:center;
  justify-content:center;font-size:20px;flex-shrink:0;
}
.tool-title{
  font-size:48px!important;font-weight:700!important;
  color:var(--on-surface)!important;letter-spacing:-0.022em!important;
  line-height:1.1!important;margin:0!important;
  font-family:var(--font)!important;
}

/* ── Scrollbar ── */
::-webkit-scrollbar{width:5px;}
::-webkit-scrollbar-track{background:var(--surface-container-low);}
::-webkit-scrollbar-thumb{background:var(--outline-variant);border-radius:3px;}

/* ── Hero ── */
.stitch-hero{
  padding:3rem 0 2rem;
  animation:fadeUp .5s ease both;
}
@keyframes fadeUp{from{opacity:0;transform:translateY(12px);}to{opacity:1;transform:translateY(0);}}
</style>
""", unsafe_allow_html=True)

# ── Constants ─────────────────────────────────────────────────────────────────
LANGUAGES = {
    "English":"English","हिन्दी (Hindi)":"Hindi","தமிழ் (Tamil)":"Tamil",
    "తెలుగు (Telugu)":"Telugu","বাংলা (Bengali)":"Bengali","मराठी (Marathi)":"Marathi",
    "ಕನ್ನಡ (Kannada)":"Kannada","മലയാളം (Malayalam)":"Malayalam",
    "ગુજરાતી (Gujarati)":"Gujarati","اردو (Urdu)":"Urdu","Español":"Spanish","Français":"French",
}
TOOL_KEYS   = ["📄 Resume Review","🎤 Mock Interview","✉️ Cover Letter","💼 LinkedIn Post","🔍 Job Decoder","🚀 Apply in One Click","🔎 Job Finder","📈 Career Path","💰 Salary Coach","📋 My History"]
TOOL_LABELS = ["📄 Resume","🎤 Interview","✉️ Cover","💼 LinkedIn","🔍 Decoder","🚀 Apply","🔎 Jobs","📈 Career","💰 Salary","📋 History"]
KEY_FILE       = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".groq_key")
ADMIN_PASSWORD = "admin@career2024"

# ── Session defaults ──────────────────────────────────────────────────────────
def _d(k,v):
    if k not in st.session_state: st.session_state[k]=v

def load_key():
    return open(KEY_FILE).read().strip() if os.path.exists(KEY_FILE) else ""

def save_key(k):
    open(KEY_FILE,"w").write(k)

_d("api_key", load_key())
_d("admin_logged_in", False)
_d("show_admin", False)
_d("active_tool", "📄 Resume Review")
_d("selected_lang", "English")
_d("interview_history", [])
_d("interview_started", False)
_d("nav_idx", 0)
_d("last_nav", TOOL_LABELS[0])
_d("history_log", [])

# ── Helpers ───────────────────────────────────────────────────────────────────
def get_client():
    k = st.session_state.get("api_key","")
    if not k:
        st.error("⚠️ No API key. Click ⚙️ top-right → Admin → enter Groq key.")
        st.stop()
    return Groq(api_key=k)

def call_ai(system, user, history=None):
    msgs = [{"role":"system","content":system}]+(history or [])+[{"role":"user","content":user}]
    with st.spinner("✨ Generating..."):
        r = get_client().chat.completions.create(model="llama-3.3-70b-versatile", max_tokens=1500, messages=msgs)
    return r.choices[0].message.content

def ln():
    l = st.session_state["selected_lang"]
    return "" if l=="English" else f"\n\nIMPORTANT: Write ENTIRE response in {l}."

def mk_docx(title, text):
    doc = Document()
    h = doc.add_heading(title,0); h.runs[0].font.color.rgb=RGBColor(0x1a,0x1a,0x2e)
    p = doc.add_paragraph(f"Generated {datetime.now().strftime('%d %B %Y')}"); p.runs[0].font.size=Pt(10)
    doc.add_paragraph()
    for line in text.split("\n"):
        if line.strip(): doc.add_paragraph(line)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def log_history(tool_name, summary, output):
    """Save result to history log"""
    entry = {
        "time": datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "tool": tool_name,
        "summary": summary[:120],
        "output": output
    }
    st.session_state["history_log"].insert(0, entry)
    if len(st.session_state["history_log"]) > 50:
        st.session_state["history_log"] = st.session_state["history_log"][:50]

def abar(key, text, title, gmail="", linkedin=False):
    st.markdown("---")
    c1,c2,c3,c4 = st.columns(4)
    with c1: st.download_button("⬇️ .docx", mk_docx(title,text), f"{key}.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document",key=f"d1_{key}")
    with c2: st.download_button("⬇️ .txt", text.encode(), f"{key}.txt","text/plain",key=f"d2_{key}")
    with c3:
        gdoc=f"https://docs.google.com/document/create?title={urllib.parse.quote(title)}"
        st.markdown(f'<a href="{gdoc}" target="_blank" style="display:inline-block;background:transparent;color:#0066cc;border:1px solid #0066cc;border-radius:980px;padding:8px 14px;border-radius:8px;font-size:13px;text-decoration:none;font-weight:500;">📄 Google Docs</a>',unsafe_allow_html=True)
    with c4:
        if gmail:
            gu=f"https://mail.google.com/mail/?view=cm&su={urllib.parse.quote(gmail)}&body={urllib.parse.quote(text[:1500])}"
            st.markdown(f'<a href="{gu}" target="_blank" style="display:inline-block;background:transparent;color:#0066cc;border:1px solid #0066cc;border-radius:980px;padding:8px 14px;border-radius:8px;font-size:13px;text-decoration:none;font-weight:500;">✉️ Gmail</a>',unsafe_allow_html=True)
        elif linkedin:
            lu=f"https://www.linkedin.com/feed/?shareActive=true&text={urllib.parse.quote(text[:700])}"
            st.markdown(f'<a href="{lu}" target="_blank" style="display:inline-block;background:transparent;color:#0066cc;border:1px solid #0066cc;border-radius:980px;padding:8px 14px;border-radius:8px;font-size:13px;text-decoration:none;font-weight:500;">🔗 LinkedIn</a>',unsafe_allow_html=True)

# ── Navigation ── Stitch Pro Digital Interface ───────────────────────────────
_TOOL_DATA = [
    ("📄 Resume Review",      "description", "Resume Review"),
    ("🎤 Mock Interview",     "forum",       "Mock Interview"),
    ("✉️ Cover Letter",       "mail",        "Cover Letter"),
    ("💼 LinkedIn Post",      "share",       "LinkedIn Post"),
    ("🔍 Job Decoder",        "analytics",   "Job Decoder"),
    ("🚀 Apply in One Click", "touch_app",   "Apply One-Click"),
    ("🔎 Job Finder",         "search",      "Job Finder"),
    ("📈 Career Path",        "map",         "Career Path"),
    ("💰 Salary Coach",       "payments",    "Salary Coach"),
    ("📋 My History",         "history",     "My History"),
]

# Read query param first
_qp = st.query_params
if "tool" in _qp:
    try:
        _ti = int(_qp["tool"])
        if 0 <= _ti < len(_TOOL_DATA):
            st.session_state["active_tool"] = _TOOL_DATA[_ti][0]
    except: pass
    st.query_params.clear()

tool_key = st.session_state["active_tool"]
_tidx = next(i for i,(tk,_,_l) in enumerate(_TOOL_DATA) if tk==tool_key)

# ── Glass Nav Bar ─────────────────────────────────────────────────────────────
st.markdown(f"""
<link href="https://fonts.googleapis.com/css2?family=Material+Symbols+Outlined:opsz,wght,FILL,GRAD@24,400,0,0&display=swap" rel="stylesheet"/>
<div class="stitch-nav">
  <div style="display:flex;align-items:center;">
    <div class="stitch-nav-brand">CareerAI</div>
    <div class="stitch-nav-links">
      <a class="stitch-nav-link" href="#">Tools</a>
      <a class="stitch-nav-link" href="#">History</a>
      <a class="stitch-nav-link" href="#">Community</a>
    </div>
  </div>
  <div class="stitch-nav-right">
    <button class="stitch-icon-btn"><span class="material-symbols-outlined">search</span></button>
    <button class="stitch-icon-btn"><span class="material-symbols-outlined">notifications</span></button>
    <button class="stitch-icon-btn"><span class="material-symbols-outlined">settings</span></button>
    <button class="stitch-upgrade-btn">Upgrade</button>
    <div class="stitch-avatar">U</div>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Left Sidebar ──────────────────────────────────────────────────────────────
_sidebar_items = "".join([
    f'<button class="stitch-nav-item{"  active" if i==_tidx else ""}" onclick="window.parent.location.href=window.parent.location.pathname+\'?tool={i}\'"  style="width:100%;text-align:left;background:none;border:none;cursor:pointer;"><span class="material-symbols-outlined ms">{icon}</span>{label}</button>'
    for i,(tk,icon,label) in enumerate(_TOOL_DATA)
])
st.markdown(f"""
<div class="stitch-sidebar">
  <div class="stitch-sidebar-header">
    <div class="stitch-sidebar-title">Career Suite</div>
    <div class="stitch-sidebar-sub">AI-Powered Tools</div>
  </div>
  {_sidebar_items}
</div>
""", unsafe_allow_html=True)

# Settings expander
with st.expander("⚙  Settings — Language & API Key", expanded=False):
    _c1, _c2, _c3 = st.columns([2,2,3])
    with _c1:
        lang_choice = st.selectbox("Language", list(LANGUAGES.keys()), key="lang_sel")
        st.session_state["selected_lang"] = LANGUAGES[lang_choice]
    with _c2:
        if not st.session_state["admin_logged_in"]:
            _pw = st.text_input("Admin password", type="password", placeholder="Password", key="adm_pw")
            if st.button("Login", key="adm_login"):
                if _pw == ADMIN_PASSWORD:
                    st.session_state["admin_logged_in"] = True; st.rerun()
                else: st.error("Wrong password")
        else:
            st.success("Logged in")
            if st.button("Logout", key="adm_out"):
                st.session_state["admin_logged_in"] = False; st.rerun()
    with _c3:
        if st.session_state["admin_logged_in"]:
            _nk = st.text_input("Groq API Key", type="password", value=st.session_state["api_key"], placeholder="gsk_...", key="adm_key")
            if st.button("Save Key", key="adm_save"):
                if _nk: st.session_state["api_key"] = _nk; save_key(_nk); st.success("Saved!")
        else:
            st.caption("Login as admin to set API key")

tool = tool_key



# ── Hero ─────────────────────────────────────────────────────────────────────
pills = " ".join([f'<span class="lang-pill">{l.split("(")[0].strip()}</span>' for l in list(LANGUAGES.keys())[:6]])
st.markdown(f"""
<div class="stitch-hero">
  <h1 class="stitch-page-title" style="font-size:48px;font-weight:700;color:var(--on-surface);letter-spacing:-0.022em;line-height:1.1;margin-bottom:8px;font-family:Inter,sans-serif;">
    Your Career,<br><span style="color:var(--primary-container);">Elevated.</span>
  </h1>
  <p style="font-size:17px;color:var(--secondary);letter-spacing:-0.022em;line-height:1.5;margin-bottom:16px;">
    Resume reviews · Mock interviews · Salary coaching · 12 languages
  </p>
  <div style="display:flex;flex-wrap:wrap;gap:0;">{pills} <span class="lang-pill">+6 more</span></div>
</div>
""", unsafe_allow_html=True)



# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Resume Review
# ─────────────────────────────────────────────────────────────────────────────
if tool == "📄 Resume Review":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">📄</div><div class="tool-title">Resume Builder</div></div>', unsafe_allow_html=True)
    st.caption("Upload or paste your resume and job description. Get a score, improved bullet points, and skill gap analysis.")
    c1,c2 = st.columns(2)
    with c1:
        st.markdown("**Your Resume**")
        up = st.file_uploader("Upload PDF or Word", type=["pdf","docx"], key="res_up")
        rtxt = ""
        if up:
            if up.type == "application/pdf":
                try:
                    import pdfplumber
                    with pdfplumber.open(up) as pdf: rtxt="\n".join([p.extract_text() or "" for p in pdf.pages])
                    st.success(f"✅ {up.name}")
                except ImportError: st.warning("pdfplumber not installed")
            else:
                try:
                    from docx import Document as D2
                    rtxt="\n".join([p.text for p in D2(up).paragraphs if p.text.strip()])
                    st.success(f"✅ {up.name}")
                except Exception as e: st.error(str(e))
        paste = st.text_area("Or paste resume text", height=180, placeholder="Paste resume here...")
        if not rtxt and paste: rtxt=paste
        if rtxt: st.markdown(f'<div style="background:rgba(201,168,76,.08);border:1px solid rgba(201,168,76,.2);border-radius:8px;padding:6px 12px;font-size:12px;color:#c9a84c;">✓ Resume ready — {len(rtxt)} chars</div>', unsafe_allow_html=True)
    with c2:
        jd = st.text_area("Job description", height=300, placeholder="Paste job posting here...")
    if st.button("✨ Review & Improve My Resume"):
        if not rtxt or not jd: st.warning("Add your resume and job description.")
        else:
            r = call_ai(f"Expert resume coach. Review resume vs job.\n1. MATCH SCORE (/10)\n2. TOP 3 IMPROVED BULLET POINTS\n3. SKILL GAPS (2-3)\n4. QUICK WINS (2 tips)\nBe specific.{ln()}", f"RESUME:\n{rtxt}\n\nJOB:\n{jd}")
            st.session_state["res_r"] = r
            log_history("📄 Resume Review", "Resume reviewed", r)
    if "res_r" in st.session_state:
        st.markdown('<div class="output-section">'+st.session_state["res_r"].replace('\n','<br>')+'</div>', unsafe_allow_html=True)
        abar("resume", st.session_state["res_r"], "Resume Review", gmail="My Reviewed Resume")

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Mock Interview
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "🎤 Mock Interview":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">🎤</div><div class="tool-title">Mock Interview</div></div>', unsafe_allow_html=True)
    st.caption("Practice real interview questions with AI feedback after every answer.")
    c1,c2,c3 = st.columns(3)
    with c1: role = st.text_input("Job role", placeholder="e.g. Data Analyst")
    with c2: lvl  = st.selectbox("Experience", ["Entry (0–2 yrs)","Mid (2–5 yrs)","Senior (5+ yrs)"])
    with c3: ityp = st.selectbox("Type", ["General / behavioral","Technical","Case study","HR round"])
    ca,cb = st.columns([1,3])
    with ca:
        if st.button("▶️ Start Interview"):
            if not role: st.warning("Enter a job role.")
            else:
                sys = f"Senior interviewer for {lvl} {role}, {ityp} interview. Ask ONE question at a time. Give 1-2 sentence feedback then next question. After 5 questions give overall score.{ln()}"
                q = call_ai(sys, "Start the interview.")
                st.session_state.update({"interview_history":[{"role":"user","content":"Start."},{"role":"assistant","content":q}],"interview_started":True,"interview_system":sys})
                st.rerun()
    with cb:
        if st.session_state["interview_started"] and st.button("🔄 Reset"):
            st.session_state.update({"interview_history":[],"interview_started":False}); st.rerun()
    if st.session_state["interview_started"] and st.session_state["interview_history"]:
        for m in st.session_state["interview_history"]:
            if m["role"]=="assistant": st.markdown(f'<div class="chat-ai">🤖 {m["content"].replace(chr(10),"<br>")}</div>', unsafe_allow_html=True)
            elif m["content"]!="Start.": st.markdown(f'<div class="chat-user">👤 {m["content"]}</div>', unsafe_allow_html=True)
        ans = st.text_input("Your answer", placeholder="Type your answer...", key="int_ans")
        if st.button("Send ↗") and ans:
            st.session_state["interview_history"].append({"role":"user","content":ans})
            r = call_ai(st.session_state["interview_system"], ans, history=st.session_state["interview_history"][:-1])
            st.session_state["interview_history"].append({"role":"assistant","content":r}); st.rerun()
        if len(st.session_state["interview_history"])>4:
            tr="\n\n".join([f"{'Interviewer' if m['role']=='assistant' else 'Candidate'}: {m['content']}" for m in st.session_state["interview_history"] if m["content"]!="Start."])
            abar("interview", tr, "Mock Interview Transcript")

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Cover Letter
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "✉️ Cover Letter":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">✉️</div><div class="tool-title">Cover Letter</div></div>', unsafe_allow_html=True)
    st.caption("Paste the job description and your details. Get a polished, human-sounding cover letter.")
    c1,c2 = st.columns(2)
    with c1:
        cjd=st.text_area("Job description", height=200, placeholder="Paste job posting...")
        cn=st.text_input("Your full name"); ce=st.text_input("Your email")
    with c2:
        ca2=st.text_area("About you", height=200, placeholder="Skills, achievements, experience...")
        ct=st.selectbox("Tone",["Professional & formal","Friendly & conversational","Confident & bold","Humble & enthusiastic"])
        cc=st.text_input("Company name (optional)")
    if st.button("✨ Generate Cover Letter"):
        if not cjd or not ca2: st.warning("Fill in job description and your details.")
        else:
            r=call_ai(f"Write a compelling {ct.lower()} cover letter. 3-4 paragraphs. Never start with 'I am writing to'. Be human and specific.{ln()}", f"Candidate:{cn}\nEmail:{ce}\nCompany:{cc or 'the company'}\nJob:\n{cjd}\nAbout:\n{ca2}")
            st.session_state["cov_r"]=r
            log_history("✉️ Cover Letter", f"Cover letter for {cc or 'a company'}", r)
    if "cov_r" in st.session_state:
        st.markdown('<div class="output-section">'+st.session_state["cov_r"].replace('\n','<br>')+'</div>', unsafe_allow_html=True)
        abar("cover", st.session_state["cov_r"], "Cover Letter", gmail=f"Application – {cc or 'Job'} – {cn}")

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: LinkedIn Post
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "💼 LinkedIn Post":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">💼</div><div class="tool-title">LinkedIn Post</div></div>', unsafe_allow_html=True)
    st.caption("Describe your achievement in rough words — get a polished post that sounds like you.")
    li=st.text_area("What do you want to post about?", height=150, placeholder="e.g. Just got promoted...")
    c1,c2=st.columns(2)
    with c1: pt=st.selectbox("Post type",["Career milestone","New job","Lesson learned","Project achievement","Thought leadership"])
    with c2: ps=st.selectbox("Style",["Storytelling & personal","Professional & concise","Inspiring","Humble & grateful"])
    if st.button("✨ Write LinkedIn Post"):
        if not li: st.warning("Describe what you want to post about.")
        else:
            r=call_ai(f"Write a {ps.lower()} LinkedIn post about {pt.lower()}. Hook in line 1. Short paragraphs. 150-250 words. 3-5 hashtags. Never start with 'Thrilled','Humbled','Excited to announce'.{ln()}", li)
            st.session_state["li_r"]=r
            log_history("💼 LinkedIn Post", f"{pt} post", r)
    if "li_r" in st.session_state:
        st.markdown('<div class="output-section">'+st.session_state["li_r"].replace('\n','<br>')+'</div>', unsafe_allow_html=True)
        abar("linkedin", st.session_state["li_r"], "LinkedIn Post", linkedin=True)

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Job Decoder
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "🔍 Job Decoder":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">🔍</div><div class="tool-title">Job Decoder</div></div>', unsafe_allow_html=True)
    st.caption("Paste any job posting. Understand what they really want, spot red flags, and know how to apply.")
    djd=st.text_area("Paste the full job description", height=300, placeholder="Paste job posting here...")
    if st.button("🔍 Decode This Job"):
        if not djd: st.warning("Paste a job description first.")
        else:
            r=call_ai(f"Decode this job posting:\n1. WHAT THEY REALLY WANT\n2. MUST-HAVE vs NICE-TO-HAVE\n3. RED FLAGS (max 3)\n4. WHAT TO HIGHLIGHT\n5. SALARY ESTIMATE\n6. 3 SMART QUESTIONS to ask{ln()}", f"Job:\n{djd}")
            st.session_state["dec_r"]=r
            log_history("🔍 Job Decoder", "Job decoded", r)
    if "dec_r" in st.session_state:
        st.markdown('<div class="output-section">'+st.session_state["dec_r"].replace('\n','<br>')+'</div>', unsafe_allow_html=True)
        abar("decoder", st.session_state["dec_r"], "Job Analysis")

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Apply in One Click
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "🚀 Apply in One Click":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">🚀</div><div class="tool-title">Apply in One Click</div></div>', unsafe_allow_html=True)
    st.caption("Get a complete application package — cover letter, email, and talking points — in one go.")
    c1,c2=st.columns(2)
    with c1:
        ajd=st.text_area("Job description", height=220, placeholder="Paste job posting...")
        aurl=st.text_input("Job URL (optional)"); ahr=st.text_input("HR email (optional)")
    with c2:
        ares=st.text_area("Your resume / key details", height=220, placeholder="Paste resume or key experience...")
        an=st.text_input("Your full name"); ae=st.text_input("Your email")
    if st.button("🚀 Generate Application Package"):
        if not ajd or not ares: st.warning("Fill in job description and your details.")
        else:
            with st.spinner("Building your complete package..."):
                cov=call_ai(f"Write a 3-paragraph cover letter. Never start with 'I am writing to'. Be specific.{ln()}", f"Candidate:{an}\nJob:\n{ajd}\nResume:\n{ares}")
                em=call_ai(f"Write a short job application email body (3-4 sentences) to send with cover letter attached.{ln()}", f"Candidate:{an}\nJob:\n{ajd[:500]}")
                pts=call_ai(f"Give 5 bullet points: strongest things this candidate should emphasise.{ln()}", f"Job:\n{ajd}\nCandidate:\n{ares}")
            st.session_state.update({"ap_cov":cov,"ap_em":em,"ap_pts":pts,"ap_jd":ajd,"ap_url":aurl,"ap_hr":ahr,"ap_name":an})
    if "ap_cov" in st.session_state:
        with st.expander("✉️ Cover Letter", expanded=True):
            st.markdown('<div class="output-section">'+st.session_state["ap_cov"].replace('\n','<br>')+'</div>', unsafe_allow_html=True)
            st.download_button("⬇️ Download .docx", mk_docx("Cover Letter",st.session_state["ap_cov"]), "cover_letter.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document",key="dl_apcov")
        with st.expander("📧 Application Email", expanded=True):
            st.markdown('<div class="output-section">'+st.session_state["ap_em"].replace('\n','<br>')+'</div>', unsafe_allow_html=True)
        with st.expander("💡 Key Talking Points", expanded=True):
            st.markdown('<div class="output-section">'+st.session_state["ap_pts"].replace('\n','<br>')+'</div>', unsafe_allow_html=True)
        st.markdown("### 🚀 Apply Now")
        st.markdown('<div class="apply-box"><h4>Quick Apply Links</h4>', unsafe_allow_html=True)
        lk=""
        if st.session_state.get("ap_url"): lk+=f'<a class="apply-link" href="{st.session_state["ap_url"]}" target="_blank">🔗 Company Site</a>'
        if st.session_state.get("ap_hr"):
            su=urllib.parse.quote(f"Job Application – {st.session_state.get('ap_name','')}"); bd=urllib.parse.quote(st.session_state["ap_em"][:1000])
            lk+=f'<a class="apply-link" href="https://mail.google.com/mail/?view=cm&to={st.session_state["ap_hr"]}&su={su}&body={bd}" target="_blank">✉️ Gmail</a>'
        kw=urllib.parse.quote(st.session_state.get("ap_jd","")[:50])
        lk+=f'<a class="apply-link" href="https://www.linkedin.com/jobs/search/?keywords={kw}" target="_blank">💼 LinkedIn</a>'
        lk+='<a class="apply-link" href="https://www.naukri.com" target="_blank">🇮🇳 Naukri</a>'
        lk+='<a class="apply-link" href="https://www.shine.com" target="_blank">✨ Shine</a>'
        st.markdown(lk+"</div>", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Job Finder
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "🔎 Job Finder":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">🔎</div><div class="tool-title">Job Finder</div></div>', unsafe_allow_html=True)
    st.caption("Upload or paste your resume. AI finds the best matching jobs from Naukri, LinkedIn, Indeed and more.")

    c1, c2 = st.columns(2)
    with c1:
        up = st.file_uploader("Upload Resume (PDF/Word)", type=["pdf","docx"], key="jf_up")
        jf_resume = ""
        if up:
            if up.type == "application/pdf":
                try:
                    import pdfplumber
                    with pdfplumber.open(up) as pdf: jf_resume = "\n".join([p.extract_text() or "" for p in pdf.pages])
                    st.success(f"✅ {up.name}")
                except: st.warning("Install pdfplumber")
            else:
                try:
                    from docx import Document as D2
                    jf_resume = "\n".join([p.text for p in D2(up).paragraphs if p.text.strip()])
                    st.success(f"✅ {up.name}")
                except Exception as e: st.error(str(e))
        jf_paste = st.text_area("Or paste resume", height=150, placeholder="Paste resume text...")
        if not jf_resume and jf_paste: jf_resume = jf_paste

    with c2:
        jf_role    = st.text_input("Job role / title", placeholder="e.g. Software Engineer, Marketing Manager")
        jf_loc     = st.text_input("Preferred location", placeholder="e.g. Mumbai, Bangalore, Remote")
        jf_exp     = st.selectbox("Experience level", ["Fresher (0-1 yr)","Junior (1-3 yrs)","Mid (3-6 yrs)","Senior (6+ yrs)"])
        jf_type    = st.selectbox("Job type", ["Full-time","Part-time","Remote","Freelance","Internship"])

    if st.button("🔎 Find Matching Jobs"):
        if not jf_resume and not jf_role:
            st.warning("Add your resume or job role to search.")
        else:
            with st.spinner("Finding best matching jobs..."):
                # AI extracts skills and suggests search terms
                skills_resp = call_ai(
                    "You are a career expert. Analyze the resume/role and extract: top 5 skills, best job titles to search for, and ideal companies. Be specific and concise.",
                    f"Resume/Role: {jf_resume[:1000] if jf_resume else jf_role}\nLocation: {jf_loc}\nExperience: {jf_exp}"
                )
                st.session_state["jf_skills"] = skills_resp
                st.session_state["jf_role_q"] = jf_role or "software engineer"
                st.session_state["jf_loc_q"] = jf_loc or "India"
                st.session_state["jf_exp_q"] = jf_exp
                st.session_state["jf_type_q"] = jf_type

    if "jf_skills" in st.session_state:
        st.markdown("### 🎯 Your Profile Analysis")
        st.markdown('<div class="output-section">' + st.session_state["jf_skills"].replace("\n","<br>") + '</div>', unsafe_allow_html=True)

        role_q = urllib.parse.quote(st.session_state["jf_role_q"])
        loc_q  = urllib.parse.quote(st.session_state["jf_loc_q"])

        st.markdown("### 🔗 Search Jobs Now")
        st.markdown('<div class="apply-box"><h4>Click to search on top job portals</h4>', unsafe_allow_html=True)
        links = f'''
<a class="apply-link" href="https://www.naukri.com/{st.session_state["jf_role_q"].lower().replace(" ","-")}-jobs-in-{st.session_state["jf_loc_q"].lower().replace(" ","-")}" target="_blank">🇮🇳 Naukri</a>
<a class="apply-link" href="https://www.linkedin.com/jobs/search/?keywords={role_q}&location={loc_q}" target="_blank">💼 LinkedIn</a>
<a class="apply-link" href="https://www.indeed.com/jobs?q={role_q}&l={loc_q}" target="_blank">🔍 Indeed</a>
<a class="apply-link" href="https://www.shine.com/job-search/{st.session_state["jf_role_q"].lower().replace(" ","-")}-jobs" target="_blank">✨ Shine</a>
<a class="apply-link" href="https://internshala.com/jobs/{st.session_state["jf_role_q"].lower().replace(" ","-")}-jobs" target="_blank">🎓 Internshala</a>
<a class="apply-link" href="https://www.glassdoor.co.in/Job/jobs.htm?sc.keyword={role_q}&locT=C&locId=0" target="_blank">🏢 Glassdoor</a>
<a class="apply-link" href="https://angel.co/jobs#find/f!%7B%22keywords%22%3A%5B%22{role_q}%22%5D%7D" target="_blank">🚀 AngelList</a>
        '''
        st.markdown(links + '</div>', unsafe_allow_html=True)

        log_history("🔎 Job Finder", f"Jobs found for {st.session_state['jf_role_q']}", st.session_state["jf_skills"])

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Career Path Advisor
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "📈 Career Path":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">📈</div><div class="tool-title">Career Path Advisor</div></div>', unsafe_allow_html=True)
    st.caption("Tell me where you are now. I'll show you exactly where to go next and how to get there.")

    c1, c2 = st.columns(2)
    with c1:
        cp_role    = st.text_input("Current job role", placeholder="e.g. Junior Software Engineer")
        cp_exp     = st.text_input("Years of experience", placeholder="e.g. 2 years")
        cp_skills  = st.text_area("Your current skills", height=120, placeholder="e.g. Python, React, SQL, team management...")
    with c2:
        cp_goal    = st.text_input("Dream job / goal (optional)", placeholder="e.g. CTO, Product Manager, Data Scientist")
        cp_domain  = st.selectbox("Industry", ["Technology","Finance","Marketing","Healthcare","Education","E-commerce","Consulting","Government","Other"])
        cp_timeline = st.selectbox("Your timeline", ["6 months","1 year","2 years","3-5 years","Long term (5+ years)"])

    if st.button("📈 Show My Career Path"):
        if not cp_role:
            st.warning("Enter your current job role.")
        else:
            r = call_ai(
                f"You are a senior career counselor. Create a detailed career path roadmap. Structure as:\n🎯 WHERE YOU ARE NOW\n🚀 NEXT STEP (6-12 months)\n📈 MID-TERM GOAL (1-3 years)\n🏆 LONG-TERM VISION (3-5 years)\nFor each step include: title, salary range (INR), skills to learn, certifications, action items.{ln()}",
                f"Current role: {cp_role}\nExperience: {cp_exp}\nSkills: {cp_skills}\nGoal: {cp_goal or 'Not specified'}\nIndustry: {cp_domain}\nTimeline: {cp_timeline}"
            )
            st.session_state["cp_r"] = r
            log_history("📈 Career Path", f"Career path for {cp_role}", r)

    if "cp_r" in st.session_state:
        st.markdown('<div class="output-section">' + st.session_state["cp_r"].replace("\n","<br>") + '</div>', unsafe_allow_html=True)
        abar("career_path", st.session_state["cp_r"], "Career Path Roadmap")

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: Salary Negotiation Coach
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "💰 Salary Coach":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">💰</div><div class="tool-title">Salary Negotiation Coach</div></div>', unsafe_allow_html=True)
    st.caption("Got an offer? Let AI coach you to negotiate the best salary. Know your worth.")

    tabs = st.tabs(["💰 Know Your Worth", "🗣️ Negotiation Script", "📧 Counter Offer Email"])

    with tabs[0]:
        st.markdown("#### What salary should you expect?")
        c1, c2 = st.columns(2)
        with c1:
            sn_role  = st.text_input("Job role", placeholder="e.g. Senior Data Analyst", key="sn_role")
            sn_exp   = st.text_input("Years of experience", placeholder="e.g. 4 years", key="sn_exp")
            sn_city  = st.text_input("City", placeholder="e.g. Bangalore, Mumbai, Delhi", key="sn_city")
        with c2:
            sn_skills = st.text_area("Your key skills", height=100, placeholder="e.g. Python, Machine Learning, SQL...", key="sn_skills")
            sn_comp   = st.text_input("Company name (optional)", placeholder="e.g. Infosys, Swiggy, Google", key="sn_comp")

        if st.button("💰 Check My Market Value", key="sn_check"):
            if not sn_role: st.warning("Enter a job role.")
            else:
                r = call_ai(
                    f"You are a salary expert for the Indian job market. Provide detailed salary analysis. Structure:\n💰 SALARY RANGE (min/mid/max in LPA)\n📊 MARKET BENCHMARK\n🏆 TOP PAYING COMPANIES\n📈 FACTORS THAT INCREASE SALARY\n💡 NEGOTIATION LEVERAGE\nBe specific with INR/LPA numbers.{ln()}",
                    f"Role: {sn_role}\nExperience: {sn_exp}\nCity: {sn_city}\nSkills: {sn_skills}\nCompany: {sn_comp or 'General market'}"
                )
                st.session_state["sn_worth"] = r
                log_history("💰 Salary Coach", f"Market value for {sn_role}", r)

        if "sn_worth" in st.session_state:
            st.markdown('<div class="output-section">' + st.session_state["sn_worth"].replace("\n","<br>") + '</div>', unsafe_allow_html=True)
            abar("salary_worth", st.session_state["sn_worth"], "Salary Analysis")

    with tabs[1]:
        st.markdown("#### Get a word-for-word negotiation script")
        c1, c2 = st.columns(2)
        with c1:
            ns_offer   = st.text_input("Offer you received (LPA)", placeholder="e.g. 8 LPA", key="ns_offer")
            ns_expect  = st.text_input("What you want (LPA)", placeholder="e.g. 12 LPA", key="ns_expect")
            ns_role2   = st.text_input("Job role", placeholder="e.g. Product Manager", key="ns_role2")
        with c2:
            ns_exp2    = st.text_input("Your experience", placeholder="e.g. 5 years", key="ns_exp2")
            ns_achieve = st.text_area("Your achievements", height=80, placeholder="e.g. Led team of 10, increased revenue by 40%...", key="ns_achieve")

        if st.button("🗣️ Write My Script", key="ns_script"):
            if not ns_offer or not ns_expect: st.warning("Enter the offer and your expected salary.")
            else:
                r = call_ai(
                    f"Write a confident, professional salary negotiation script. Include:\n📞 OPENING LINE\n💪 VALUE STATEMENT (why you deserve more)\n💬 EXACT WORDS TO SAY\n🔄 HANDLING OBJECTIONS\n✅ CLOSING\nMake it natural and confident, not aggressive.{ln()}",
                    f"Offer: {ns_offer} LPA\nExpecting: {ns_expect} LPA\nRole: {ns_role2}\nExperience: {ns_exp2}\nAchievements: {ns_achieve}"
                )
                st.session_state["ns_script_r"] = r
                log_history("💰 Negotiation Script", f"Script for {ns_role2}", r)

        if "ns_script_r" in st.session_state:
            st.markdown('<div class="output-section">' + st.session_state["ns_script_r"].replace("\n","<br>") + '</div>', unsafe_allow_html=True)
            abar("neg_script", st.session_state["ns_script_r"], "Negotiation Script")

    with tabs[2]:
        st.markdown("#### Write a counter offer email")
        c1, c2 = st.columns(2)
        with c1:
            ce_name    = st.text_input("Your name", key="ce_name")
            ce_offer   = st.text_input("Offer received (LPA)", key="ce_offer")
            ce_expect  = st.text_input("Counter offer (LPA)", key="ce_expect")
        with c2:
            ce_role    = st.text_input("Role", key="ce_role")
            ce_company = st.text_input("Company", key="ce_company")
            ce_reason  = st.text_area("Why you deserve more", height=80, placeholder="e.g. 5 years experience, led major projects...", key="ce_reason")

        if st.button("📧 Write Counter Offer Email", key="ce_btn"):
            if not ce_offer or not ce_expect: st.warning("Enter offer and counter offer amounts.")
            else:
                r = call_ai(
                    f"Write a professional, confident counter offer email. Be polite but firm. Show value. 3-4 paragraphs.{ln()}",
                    f"Name: {ce_name}\nCompany: {ce_company}\nRole: {ce_role}\nOffer: {ce_offer} LPA\nCounter: {ce_expect} LPA\nReason: {ce_reason}"
                )
                st.session_state["ce_r"] = r
                log_history("💰 Counter Offer", f"Counter offer for {ce_company}", r)

        if "ce_r" in st.session_state:
            st.markdown('<div class="output-section">' + st.session_state["ce_r"].replace("\n","<br>") + '</div>', unsafe_allow_html=True)
            abar("counter_email", st.session_state["ce_r"], "Counter Offer Email", gmail=f"Counter Offer – {ce_role} – {ce_company}")

# ─────────────────────────────────────────────────────────────────────────────
# TOOL: My History
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "📋 My History":
    st.markdown('<div style="margin-bottom:8px;"><div style="display:flex;align-items:center;gap:12px;margin-bottom:4px;"><div class="tool-icon">📋</div><div class="tool-title">My History</div></div>', unsafe_allow_html=True)
    st.caption("All your generated content saved in one place. Download or reuse anytime.")

    history = st.session_state.get("history_log", [])

    if not history:
        st.markdown('''
<div style="text-align:center;padding:3rem;color:rgba(0,0,0,.48);">
    <div style="font-size:3rem;margin-bottom:1rem;">📋</div>
    <div style="font-size:1.1rem;font-weight:500;margin-bottom:.5rem;">No history yet</div>
    <div style="font-size:.9rem;">Use any tool above and your results will be saved here automatically</div>
</div>
''', unsafe_allow_html=True)
    else:
        c1, c2 = st.columns([3,1])
        with c1: st.markdown(f"**{len(history)} saved items**")
        with c2:
            if st.button("🗑️ Clear History", key="clear_hist"):
                st.session_state["history_log"] = []
                st.rerun()

        st.markdown("---")
        for i, entry in enumerate(history):
            with st.expander(f"{entry['tool']} · {entry['time']} — {entry['summary']}...", expanded=(i==0)):
                st.markdown('<div class="output-section">' + entry["output"].replace("\n","<br>") + '</div>', unsafe_allow_html=True)
                c1,c2,c3 = st.columns(3)
                with c1: st.download_button("⬇️ .docx", mk_docx(entry["tool"], entry["output"]), f"history_{i}.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document",key=f"hd1_{i}")
                with c2: st.download_button("⬇️ .txt", entry["output"].encode(), f"history_{i}.txt","text/plain",key=f"hd2_{i}")
                with c3:
                    gdoc=f"https://docs.google.com/document/create?title={urllib.parse.quote(entry['tool'])}"
                    st.markdown(f'<a href="{gdoc}" target="_blank" style="display:inline-block;background:transparent;color:#0066cc;border:1px solid #0066cc;border-radius:980px;padding:8px 14px;border-radius:8px;font-size:13px;text-decoration:none;font-weight:500;">📄 Google Docs</a>',unsafe_allow_html=True)
